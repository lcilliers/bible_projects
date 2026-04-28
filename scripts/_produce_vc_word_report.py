"""
Produce a Verse Context word report — shows the full classification result
for a single registry after Verse Context processing.

Organised by Strong's root number, with sub-glosses (suffixed variants) nested.
Each sub-gloss shows its contextual meaning groups, anchor verses, related verses,
and set-aside verses.

Usage:
    python scripts/_produce_vc_word_report.py --registry=4
    python scripts/_produce_vc_word_report.py --registry=4 --include-deleted
"""
import sqlite3
import argparse
import os
import re
from datetime import date
from collections import OrderedDict

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("WARNING: python-docx not installed. Producing markdown instead.")

DB_PATH = os.path.join('database', 'bible_research.db')


def strip_suffix(strongs):
    """H2734A -> H2734, G3709 -> G3709."""
    if strongs:
        return re.match(r'^[HG]\d+', strongs).group()
    return strongs


def build_report_data(conn, registry_no, include_deleted=False):
    """Gather all data organised by Strong's root.

    Args:
        include_deleted: If True, include delete_flagged terms and verses.
                        If False (default), exclude them.
    """
    reg = conn.execute("SELECT * FROM word_registry WHERE no = ?", (registry_no,)).fetchone()
    if not reg:
        raise ValueError(f"Registry {registry_no} not found")

    del_filter_ti = "" if include_deleted else "AND ti.delete_flagged = 0"
    del_filter_mt = "" if include_deleted else "AND mt.delete_flagged = 0"

    # OWNER + XREF terms
    terms = conn.execute(f'''
        SELECT ti.strongs_number, ti.transliteration, ti.step_search_gloss,
               ti.meaning, ti.term_owner_type, ti.occurrence_count,
               ti.delete_flagged as ti_deleted,
               mt.id as mti_id, mt.gloss as mti_gloss, mt.status as mti_status
        FROM wa_term_inventory ti
        JOIN wa_file_index fi ON fi.id = ti.file_id
        JOIN word_registry wr ON wr.id = fi.word_registry_fk
        LEFT JOIN mti_terms mt ON mt.strongs_number = ti.strongs_number {del_filter_mt}
        WHERE wr.no = ? {del_filter_ti}
        ORDER BY ti.strongs_number
    ''', (registry_no,)).fetchall()

    # Group terms by base Strong's number
    roots = OrderedDict()
    for t in terms:
        base = strip_suffix(t['strongs_number'])
        if base not in roots:
            roots[base] = {
                'base_strongs': base,
                'transliteration': t['transliteration'],
                'sub_glosses': []
            }

        del_filter_vr = "" if include_deleted else "AND vr.delete_flagged = 0"
        del_filter_vc = "" if include_deleted else "AND vc.delete_flagged = 0"
        del_filter_vcg = "" if include_deleted else "AND delete_flagged = 0"

        entry = {
            'strongs_number': t['strongs_number'],
            'gloss': t['step_search_gloss'] or t['mti_gloss'] or '',
            'meaning': t['meaning'],
            'term_owner_type': t['term_owner_type'],
            'occurrence_count': t['occurrence_count'],
            'mti_status': t['mti_status'],
            'mti_id': t['mti_id'],
            'ti_deleted': t['ti_deleted'],
            'groups': [],
            'set_aside': [],
            'senses': [],
            'stems': [],
            'lsj': None
        }

        # Meaning parse data
        ti_id = conn.execute(
            f'''SELECT ti.id FROM wa_term_inventory ti
               JOIN wa_file_index fi ON fi.id = ti.file_id
               WHERE fi.word_registry_fk = (SELECT id FROM word_registry WHERE no = ?)
                 AND ti.strongs_number = ? {del_filter_ti}''',
            (registry_no, t['strongs_number'])
        ).fetchone()
        if ti_id:
            senses = conn.execute('''
                SELECT ms.level_code, ms.sense_text, ms.stem_label, ms.domain_tag
                FROM wa_meaning_sense ms
                JOIN wa_meaning_parsed mp ON mp.id = ms.parsed_meaning_id
                WHERE mp.term_inv_id = ?
                ORDER BY ms.sort_order
            ''', (ti_id['id'],)).fetchall()
            entry['senses'] = [{'code': s['level_code'], 'text': s['sense_text'],
                                'stem': s['stem_label'], 'domain': s['domain_tag']} for s in senses]

            stems = conn.execute('''
                SELECT st.stem_name, st.stem_type, st.top_sense_text, st.sense_count
                FROM wa_meaning_stem st
                JOIN wa_meaning_parsed mp ON mp.id = st.parsed_meaning_id
                WHERE mp.term_inv_id = ?
            ''', (ti_id['id'],)).fetchall()
            entry['stems'] = [{'name': s['stem_name'], 'type': s['stem_type'],
                               'sense': s['top_sense_text'], 'count': s['sense_count']} for s in stems]

            lsj = conn.execute(
                'SELECT lsj_gloss, lsj_domains, lsj_philosophical_note, lsj_etymology_note FROM wa_lsj_parsed WHERE term_inv_id = ?',
                (ti_id['id'],)
            ).fetchone()
            if lsj and lsj['lsj_gloss']:
                entry['lsj'] = {'gloss': lsj['lsj_gloss'], 'domains': lsj['lsj_domains'],
                                'philosophy': lsj['lsj_philosophical_note'],
                                'etymology': lsj['lsj_etymology_note']}

        # Get groups and verses for OWNER terms with verse_context
        if t['mti_id'] and t['term_owner_type'] == 'OWNER':
            groups = conn.execute(f'''
                SELECT id, group_code, context_description, notes
                FROM verse_context_group
                WHERE mti_term_id = ? {del_filter_vcg}
                ORDER BY group_code
            ''', (t['mti_id'],)).fetchall()

            for g in groups:
                gdata = {
                    'group_code': g['group_code'],
                    'description': g['context_description'],
                    'notes': g['notes'],
                    'anchors': [],
                    'related': []
                }
                for a in conn.execute(f'''
                    SELECT vr.reference, vr.verse_text
                    FROM verse_context vc
                    JOIN wa_verse_records vr ON vr.id = vc.verse_record_id
                    WHERE vc.group_id = ? AND vc.is_anchor = 1
                      {del_filter_vc} {del_filter_vr}
                    ORDER BY vr.book_id, vr.chapter, vr.verse_num
                ''', (g['id'],)).fetchall():
                    gdata['anchors'].append({'ref': a['reference'], 'text': a['verse_text']})

                for r in conn.execute(f'''
                    SELECT vr.reference, vr.verse_text
                    FROM verse_context vc
                    JOIN wa_verse_records vr ON vr.id = vc.verse_record_id
                    WHERE vc.group_id = ? AND vc.is_related = 1
                      {del_filter_vc} {del_filter_vr}
                    ORDER BY vr.book_id, vr.chapter, vr.verse_num
                ''', (g['id'],)).fetchall():
                    gdata['related'].append({'ref': r['reference'], 'text': r['verse_text']})

                entry['groups'].append(gdata)

            # Set aside
            for s in conn.execute(f'''
                SELECT vr.reference, vr.verse_text
                FROM verse_context vc
                JOIN wa_verse_records vr ON vr.id = vc.verse_record_id
                WHERE vc.mti_term_id = ? AND vc.is_relevant = 0
                  {del_filter_vc} {del_filter_vr}
                ORDER BY vr.book_id, vr.chapter, vr.verse_num
            ''', (t['mti_id'],)).fetchall():
                entry['set_aside'].append({'ref': s['reference'], 'text': s['verse_text']})

        roots[base]['sub_glosses'].append(entry)

    return reg, roots


def produce_docx(reg, roots, outpath, include_deleted=False):
    """Produce .docx report."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(10)

    doc.add_heading(f"Verse Context Report: {reg['word']} (Registry {reg['no']})", level=0)
    doc.add_paragraph(
        f"Status: {reg['verse_context_status']}  |  "
        f"Cluster: {reg['cluster_assignment']}  |  "
        f"Produced: {date.today().isoformat()}  |  "
        f"{'Includes deleted records' if include_deleted else 'Excludes deleted records'}"
    )

    # Summary counts
    total_owner = sum(1 for r in roots.values() for sg in r['sub_glosses'] if sg['term_owner_type'] == 'OWNER')
    total_xref = sum(1 for r in roots.values() for sg in r['sub_glosses'] if sg['term_owner_type'] == 'XREF')
    total_groups = sum(len(sg['groups']) for r in roots.values() for sg in r['sub_glosses'])
    total_anchors = sum(len(g['anchors']) for r in roots.values() for sg in r['sub_glosses'] for g in sg['groups'])
    total_related = sum(len(g['related']) for r in roots.values() for sg in r['sub_glosses'] for g in sg['groups'])
    total_aside = sum(len(sg['set_aside']) for r in roots.values() for sg in r['sub_glosses'])

    p = doc.add_paragraph()
    p.add_run(f"OWNER terms: {total_owner}  |  XREF terms: {total_xref}  |  "
              f"Groups: {total_groups}  |  Anchors: {total_anchors}  |  "
              f"Related: {total_related}  |  Set aside: {total_aside}")
    doc.add_paragraph("")

    for base, root in roots.items():
        # Root heading
        doc.add_heading(f"{base} — {root['transliteration']}", level=1)

        for sg in root['sub_glosses']:
            # Sub-gloss heading
            suffix = sg['strongs_number'][len(base):] if len(sg['strongs_number']) > len(base) else ''
            label = f"{sg['strongs_number']}"
            if suffix:
                label += f" (sub-gloss: {sg['gloss']})"
            else:
                label += f" — {sg['gloss']}"

            type_label = f"[{sg['term_owner_type']}]"
            status_label = f"status: {sg['mti_status'] or 'NULL'}"
            occ_label = f"occ: {sg['occurrence_count'] or '?'}"

            doc.add_heading(f"{label}  {type_label}", level=2)
            meta = doc.add_paragraph()
            meta.add_run(f"{status_label}  |  {occ_label}").italic = True

            if sg['meaning']:
                mp = doc.add_paragraph()
                mp.add_run("Meaning (raw): ").bold = True
                mp.add_run(sg['meaning'][:500])

            # Parsed senses
            if sg['senses']:
                doc.add_paragraph().add_run("Parsed senses:").bold = True
                for s in sg['senses']:
                    bp = doc.add_paragraph(style='List Bullet')
                    label = s['code'] or ''
                    text = s['text'] or ''
                    if s['stem']:
                        text += f"  [{s['stem']}]"
                    if s['domain']:
                        text += f"  ({s['domain']})"
                    bp.add_run(f"{label}  ").bold = True
                    bp.add_run(text)
                    for run in bp.runs:
                        run.font.size = Pt(9)

            # Stems
            if sg['stems']:
                doc.add_paragraph().add_run("Stems:").bold = True
                for s in sg['stems']:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.add_run(f"{s['name']}").bold = True
                    bp.add_run(f" ({s['type']}): {s['sense']}  [{s['count']} senses]")
                    for run in bp.runs:
                        run.font.size = Pt(9)

            # LSJ
            if sg['lsj']:
                lp = doc.add_paragraph()
                lp.add_run("LSJ: ").bold = True
                lp.add_run(sg['lsj']['gloss'] or '')
                if sg['lsj']['domains']:
                    lp.add_run(f"  Domains: {sg['lsj']['domains']}")
                if sg['lsj']['etymology']:
                    ep = doc.add_paragraph()
                    ep.add_run("Etymology: ").bold = True
                    ep.add_run(sg['lsj']['etymology'])

            if sg['term_owner_type'] == 'XREF':
                doc.add_paragraph("Cross-reference term — verse context derived from OWNER registry.").italic = True
                continue

            if not sg['groups'] and not sg['set_aside']:
                doc.add_paragraph("No verse context records.").italic = True
                continue

            for g in sg['groups']:
                doc.add_heading(f"Group {g['group_code']}: {g['description']}", level=3)
                if g['notes']:
                    np = doc.add_paragraph()
                    np.add_run("Note: ").bold = True
                    np.add_run(g['notes']).italic = True

                if g['anchors']:
                    doc.add_paragraph().add_run("Anchor verses:").bold = True
                    for a in g['anchors']:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.add_run(f"{a['ref']}").bold = True
                        bp.add_run(f" — {a['text']}")
                        for run in bp.runs:
                            run.font.size = Pt(9)

                if g['related']:
                    doc.add_paragraph().add_run(f"Related verses ({len(g['related'])}):").bold = True
                    for r in g['related']:
                        bp = doc.add_paragraph(style='List Bullet')
                        bp.add_run(f"{r['ref']}").bold = True
                        bp.add_run(f" — {r['text']}")
                        for run in bp.runs:
                            run.font.size = Pt(9)

            if sg['set_aside']:
                doc.add_heading(f"Set Aside ({len(sg['set_aside'])} verses)", level=3)
                for s in sg['set_aside']:
                    bp = doc.add_paragraph(style='List Bullet')
                    bp.add_run(f"{s['ref']}").bold = True
                    bp.add_run(f" — {s['text']}")
                    for run in bp.runs:
                        run.font.size = Pt(9)

    doc.save(outpath)
    print(f"Written: {outpath}")


def main():
    parser = argparse.ArgumentParser(description='Produce Verse Context word report')
    parser.add_argument('--registry', type=int, required=True)
    parser.add_argument('--include-deleted', action='store_true',
                        help='Include delete_flagged terms and verses')
    args = parser.parse_args()

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row

    reg, roots = build_report_data(conn, args.registry, include_deleted=args.include_deleted)
    word = reg['word']
    reg_no = reg['no']
    today_c = date.today().strftime('%Y%m%d')

    if HAS_DOCX:
        outpath = os.path.join('outputs', 'reports', 'words',
                               f'vc-report-{reg_no:03d}-{word}-{today_c}.docx')
        produce_docx(reg, roots, outpath, include_deleted=args.include_deleted)
    else:
        outpath = os.path.join('outputs', 'reports', 'words',
                               f'vc-report-{reg_no:03d}-{word}-{today_c}.md')
        print(f"Would write markdown to {outpath}")

    conn.close()


if __name__ == '__main__':
    main()
