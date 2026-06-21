"""Render a single integrated-essay markdown file to .docx (v2).

Successor to _render_essay_to_docx_v1_20260527.py. Adds:
  - tolerates an optional bullet prefix ("- ", "* ", "+ ") on admin front-matter
    lines (the cluster-essay header carries File/Date/Version + a provenance line
    as bullets), so the admin block is stripped from the reader-facing docx;
  - renders genuine bullet lists with the Word 'List Bullet' style;
  - everything v1 did: #/##/### headings, **bold**, *italic*, --- rules,
    > blockquotes, plain paragraphs.

Usage:
    python scripts/_render_essay_to_docx_v2_20260620.py --in <essay.md> [--out <file.docx>]

Output: .docx sibling to the .md input unless --out is given.
"""
from __future__ import annotations
import argparse, io, re, sys
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches

RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")
RE_HR = re.compile(r"^---+\s*$")
RE_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
RE_BULLET = re.compile(r"^\s*[-*+]\s+(.*)$")
# admin front-matter (after an optional bullet marker is stripped): the metadata
# line(s) and the italic provenance line — never reader content.
RE_FRONT_MATTER = re.compile(
    r"^(?:\*\*(?:File|Programme reference|Authored|Cluster|Date|Version|Tier|Source|Status):\*\*"
    r"|\*Written to\b)", re.IGNORECASE)


def _debullet(s: str) -> tuple[str, bool]:
    m = RE_BULLET.match(s)
    return (m.group(1).strip(), True) if m else (s, False)


def add_inline_runs(para, text: str) -> None:
    cursor = 0
    while cursor < len(text):
        nb = RE_BOLD.search(text, cursor)
        ni = RE_ITALIC.search(text, cursor)
        cands = [m for m in (nb, ni) if m is not None]
        if not cands:
            para.add_run(text[cursor:])
            break
        m = min(cands, key=lambda x: x.start())
        if m.start() > cursor:
            para.add_run(text[cursor:m.start()])
        run = para.add_run(m.group(1))
        if m.re is RE_BOLD:
            run.bold = True
        else:
            run.italic = True
        cursor = m.end()


def render(in_path: Path, out_path: Path) -> None:
    lines = in_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # admin front-matter, with or without a leading bullet marker
        content, _ = _debullet(stripped)
        if RE_FRONT_MATTER.match(content):
            i += 1
            continue

        if RE_HR.match(stripped):
            doc.add_paragraph()
            i += 1
            continue

        m = RE_HEADING.match(stripped)
        if m:
            txt = RE_ITALIC.sub(r"\1", RE_BOLD.sub(r"\1", m.group(2)))
            doc.add_heading(txt, level=min(len(m.group(1)), 4))
            i += 1
            continue

        if stripped.startswith(">"):
            ql: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                ql.append(re.sub(r"^\s*>\s?", "", lines[i]).rstrip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            add_inline_runs(p, " ".join(q for q in ql if q.strip()))
            for r in p.runs:
                r.italic = True
            continue

        # genuine bullet list item (front-matter bullets already consumed above)
        bm = RE_BULLET.match(line)
        if bm:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, bm.group(1).strip())
            i += 1
            continue

        # plain paragraph: accumulate consecutive non-special lines
        para_lines = [line]
        i += 1
        while i < len(lines):
            ns = lines[i].strip()
            if not ns or RE_HEADING.match(ns) or RE_HR.match(ns) or ns.startswith(">") or RE_BULLET.match(ns):
                break
            if RE_FRONT_MATTER.match(_debullet(ns)[0]):
                break
            para_lines.append(lines[i])
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(l.strip() for l in para_lines))

    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_path", required=True, type=Path)
    ap.add_argument("--out", dest="out_path", default=None, type=Path)
    args = ap.parse_args()
    if not args.in_path.exists():
        print(f"FAIL: input not found: {args.in_path}")
        return 1
    render(args.in_path, args.out_path or args.in_path.with_suffix(".docx"))
    return 0


if __name__ == "__main__":
    sys.exit(main())
