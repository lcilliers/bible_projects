"""build_programme_prose_extract.py — Programme-stage prose extract (M34).

Source of truth: `prose_section_type` where `source_stage='programme'` + the
actual `prose_section` content for each type (where populated).

Usage:
  python scripts/build_programme_prose_extract.py
  python scripts/build_programme_prose_extract.py --also-markdown
  python scripts/build_programme_prose_extract.py --also-docx       # readable Word doc in outputs/docx/
  python scripts/build_programme_prose_extract.py --include-body    # include full prose body text in JSON
  python scripts/build_programme_prose_extract.py --all-formats     # JSON + MD + DOCX with bodies

Default outputs:
  data/exports/reference/wa-programme-prose-extract-{YYYYMMDD}.json
  data/exports/reference/wa-programme-prose-extract-{YYYYMMDD}.md   (if --also-markdown)
  outputs/docx/wa-programme-prose-extract-{YYYYMMDD}.docx            (if --also-docx)

Content is expected to be empty after M34 seed; populated later via PROSE
patches as researcher + Claude AI draft each programme-stage narrative.
"""
from __future__ import annotations

import argparse
import json
import os
import sqlite3
import sys
from datetime import datetime, timezone
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

DB_PATH = os.path.join("data", "bible_research.db")
OUT_DIR = os.path.join("data", "exports", "reference")
DOCX_OUT_DIR = os.path.join("outputs", "docx")
EXTRACTOR_VERSION = "1.1"

# Chapter name map (aligned to programme-prose-structure-design-v1)
CHAPTER_NAMES = {
    0: "Preamble",
    1: "Programme purpose",
    2: "Research methodology",
    3: "Research approach",
    4: "Data architecture",
    5: "Data integrity & governance",
    6: "Instruction corpus",
}


def open_db(path: str = DB_PATH) -> sqlite3.Connection:
    conn = sqlite3.connect(path)
    conn.row_factory = sqlite3.Row
    return conn


def now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def today_compact() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%d")


def get_schema_version(conn) -> str:
    row = conn.execute(
        "SELECT version_code FROM schema_version WHERE id = (SELECT MAX(id) FROM schema_version)"
    ).fetchone()
    return row[0] if row else "unknown"


def extract_programme_prose(conn, include_body: bool = False) -> dict:
    # Section type catalogue
    type_rows = conn.execute(
        """SELECT id, code, label, description, chapter_no, lifecycle_tag,
                  expected_length_min, expected_length_max, sort_order
             FROM prose_section_type
            WHERE source_stage = 'programme' AND delete_flagged = 0
            ORDER BY sort_order, code"""
    ).fetchall()

    types: list[dict] = []
    type_count = 0
    section_total = 0
    for t in type_rows:
        type_count += 1
        # Count content sections for this type
        sections = conn.execute(
            """SELECT id, registry_id, heading, status, version, author,
                      word_count, created_at, approved_at
                 FROM prose_section
                WHERE section_type_id = ? AND delete_flagged = 0
                ORDER BY version DESC""",
            (t["id"],),
        ).fetchall()
        entry = {
            "id": t["id"],  # prose_section_type.id — referenced as section_type_id in PROSE patch inserts
            "code": t["code"],
            "label": t["label"],
            "description": t["description"],
            "chapter_no": t["chapter_no"],
            "lifecycle_tag": t["lifecycle_tag"],
            "expected_length_min": t["expected_length_min"],
            "expected_length_max": t["expected_length_max"],
            "sort_order": t["sort_order"],
            "section_count": len(sections),
            "sections_preview": [dict(s) for s in sections],
        }
        if include_body and sections:
            bodies = {}
            for s in sections:
                body_row = conn.execute(
                    "SELECT body FROM prose_section WHERE id = ?", (s["id"],)
                ).fetchone()
                bodies[s["id"]] = body_row["body"] if body_row else None
            entry["bodies_by_id"] = bodies
        types.append(entry)
        section_total += len(sections)

    return {
        "types": types,
        "type_count": type_count,
        "section_total": section_total,
    }


def build_extract(conn, include_body: bool = False) -> dict:
    pp = extract_programme_prose(conn, include_body=include_body)
    return {
        "meta": {
            "generated_at": now_iso(),
            "schema_version": get_schema_version(conn),
            "extractor_version": EXTRACTOR_VERSION,
            "source": "prose_section_type WHERE source_stage='programme' (M34 live, 2026-04-20)",
            "canonical_note": "DB is source of truth post-M34 for programme-stage narrative (L2 of 3-layer reference).",
            "include_body": include_body,
            "description": "Programme-wide narrative section index — anchor verse definition, XREF architecture, validation standard, etc. Content populated via PROSE patches as researcher + AI draft each narrative.",
            "patch_hint": "To insert programme-wide prose via a PROSE patch, use `section_type_id = <id>` from each entry in programme_prose.types. Pair with `registry_id = null` (requires the schema enablement directive per wa-directive-instruction [current] §10 to relax the NOT NULL constraint).",
        },
        "programme_prose": pp,
    }


def render_markdown_view(extract: dict) -> str:
    lines = []
    meta = extract["meta"]
    pp = extract["programme_prose"]

    lines.append(f"# WA Programme-Stage Prose — {meta['generated_at'][:10]}\n")
    lines.append(f"_Schema {meta['schema_version']} · source: `prose_section_type` + `prose_section`._\n")
    lines.append("---\n")
    lines.append("## Summary\n")
    lines.append(f"**Section types seeded:** {pp['type_count']}  ·  "
                 f"**Content sections populated:** {pp['section_total']}\n")
    lines.append("---\n")

    # Partition types into (a) chaptered types that have at least one populated
    # section, for rendering as readable prose grouped by chapter; (b) all
    # remaining types (stubs not yet populated, or legacy/unchaptered), rendered
    # as a flat index at the end.
    populated = [t for t in pp["types"] if t["section_count"] > 0]
    stubs = [t for t in pp["types"] if t["section_count"] == 0]

    if populated:
        lines.append("## Programme\n")
        # Group populated types by chapter_no (None sorts last)
        by_chapter: dict = {}
        for t in populated:
            ch = t.get("chapter_no")
            by_chapter.setdefault(ch, []).append(t)
        chapter_keys = sorted(by_chapter.keys(), key=lambda c: (c is None, c if c is not None else 0))

        for ch in chapter_keys:
            chapter_name = CHAPTER_NAMES.get(ch, "Unchaptered") if ch is not None else "Unchaptered"
            header = f"Chapter {ch} — {chapter_name}" if ch is not None else chapter_name
            lines.append(f"### {header}\n")
            for t in sorted(by_chapter[ch], key=lambda x: (x["sort_order"] or 0)):
                lines.append(f"#### {t['label']}\n")
                lines.append(f"_`{t['code']}`  ·  type id {t['id']}  ·  sort {t['sort_order']}_\n")
                if t.get("description"):
                    lines.append(f"> {t['description']}\n")
                # Render each populated section's body inline (highest-version first
                # per the underlying query in extract_programme_prose).
                bodies = t.get("bodies_by_id") or {}
                for s in t["sections_preview"]:
                    body = bodies.get(s["id"])
                    if body is None:
                        # No body (either --include-body not set, or row missing). Show metadata.
                        lines.append(f"*(metadata only — body not included in this extract. "
                                     f"Run `build_programme_prose_extract.py --include-body` to render full text.)*")
                        lines.append(f"- Section id {s['id']} · status `{s['status']}` · "
                                     f"v{s['version']} · {s['word_count']} words · author `{s['author']}`\n")
                    else:
                        meta_line = (f"*Section id {s['id']} · status `{s['status']}` · "
                                     f"v{s['version']} · {s['word_count']} words · author `{s['author']}`*")
                        lines.append(meta_line + "\n")
                        lines.append(body.rstrip())
                        lines.append("")  # blank line after body
            lines.append("")  # blank line between chapters

    if stubs:
        lines.append("---\n")
        lines.append("## Section types not yet populated\n")
        lines.append("_Stubs — `prose_section_type` rows with `chapter_no=NULL` or no `prose_section` content. "
                     "`id` is the value to use as `section_type_id` in a PROSE patch insert._\n")
        lines.append("| id | code | label | chapter | sort | description |")
        lines.append("|---:|---|---|---:|---:|---|")
        for t in sorted(stubs, key=lambda x: (x["chapter_no"] is None, x["chapter_no"] or 0, x["sort_order"] or 0)):
            desc = (t.get("description") or "").replace("|", "\\|").replace("\n", " ")
            lines.append(f"| {t['id']} | `{t['code']}` | {t['label']} | "
                         f"{t['chapter_no'] if t['chapter_no'] is not None else '—'} | "
                         f"{t['sort_order']} | {desc} |")
        lines.append("")

    lines.append("---")
    lines.append(f"*Generated {meta['generated_at']}.*")
    return "\n".join(lines)


def render_docx(extract: dict, out_path: Path) -> None:
    """Readable .docx export for reading on a second monitor.

    Structure:
      Title (Heading 1): "Programme Prose"
      Subtitle: schema version + generated date
      ## Programme (Heading 1)
        ### Chapter N — {chapter name} (Heading 2, per populated chapter)
          #### {section label} (Heading 3)
            meta line (italic)
            stub description (italic)
            body paragraphs (normal, split on double-newline)
      ## Section types not yet populated (Heading 1)
        table of stubs
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[WARN] python-docx not installed; skipping .docx export")
        return

    meta = extract["meta"]
    pp = extract["programme_prose"]

    doc = Document()

    # Baseline font tweak — make body readable on a reading monitor
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Programme Prose", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph()
    r = subtitle.add_run(
        f"Schema {meta['schema_version']} · extractor v{meta['extractor_version']} · "
        f"generated {meta['generated_at']}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    summary = doc.add_paragraph()
    r = summary.add_run(
        f"Section types seeded: {pp['type_count']}  ·  "
        f"Content sections populated: {pp['section_total']}"
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    populated = [t for t in pp["types"] if t["section_count"] > 0]
    stubs = [t for t in pp["types"] if t["section_count"] == 0]

    if populated:
        doc.add_heading("Programme", level=1)

        by_chapter: dict = {}
        for t in populated:
            ch = t.get("chapter_no")
            by_chapter.setdefault(ch, []).append(t)
        chapter_keys = sorted(by_chapter.keys(), key=lambda c: (c is None, c if c is not None else 0))

        for ch in chapter_keys:
            chapter_name = CHAPTER_NAMES.get(ch, "Unchaptered") if ch is not None else "Unchaptered"
            header = f"Chapter {ch} — {chapter_name}" if ch is not None else chapter_name
            doc.add_heading(header, level=2)

            for t in sorted(by_chapter[ch], key=lambda x: (x["sort_order"] or 0)):
                doc.add_heading(t["label"], level=3)

                # Code + metadata line (small grey)
                meta_p = doc.add_paragraph()
                rr = meta_p.add_run(
                    f"code: {t['code']}  ·  type id {t['id']}  ·  sort {t['sort_order']}"
                )
                rr.italic = True
                rr.font.size = Pt(9)
                rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                # Stub description (blockquote-style — slightly indented, italic)
                if t.get("description"):
                    desc_p = doc.add_paragraph(style="Intense Quote")
                    desc_p.add_run(t["description"])

                bodies = t.get("bodies_by_id") or {}
                for s in t["sections_preview"]:
                    sec_meta = doc.add_paragraph()
                    rr = sec_meta.add_run(
                        f"Section id {s['id']}  ·  status {s['status']}  ·  v{s['version']}  ·  "
                        f"{s['word_count']} words  ·  author {s['author']}"
                    )
                    rr.italic = True
                    rr.font.size = Pt(9)
                    rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

                    body = bodies.get(s["id"])
                    if body is None:
                        p = doc.add_paragraph()
                        rr = p.add_run("(body not included in this extract — re-run with --include-body)")
                        rr.italic = True
                        continue
                    # Split paragraphs on blank lines and render each as its own paragraph
                    for para in body.split("\n\n"):
                        para = para.strip()
                        if not para:
                            continue
                        doc.add_paragraph(para)

    if stubs:
        doc.add_heading("Section types not yet populated", level=1)
        caption = doc.add_paragraph()
        rr = caption.add_run(
            "Stubs — prose_section_type rows with chapter_no=NULL or no prose_section content. "
            "id is the value to use as section_type_id in a PROSE patch insert."
        )
        rr.italic = True
        rr.font.size = Pt(9)
        rr.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "code"
        hdr[2].text = "label"
        hdr[3].text = "chapter"
        hdr[4].text = "description"
        for t in sorted(stubs, key=lambda x: (x["chapter_no"] is None, x["chapter_no"] or 0, x["sort_order"] or 0)):
            row = table.add_row().cells
            row[0].text = str(t["id"])
            row[1].text = t["code"]
            row[2].text = t["label"]
            row[3].text = str(t["chapter_no"]) if t["chapter_no"] is not None else "—"
            row[4].text = (t.get("description") or "").replace("\n", " ")

    doc.save(out_path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Build programme-stage prose extract from DB")
    ap.add_argument("--db", type=str, default=DB_PATH)
    ap.add_argument("--out", type=str, default=None)
    ap.add_argument("--also-markdown", action="store_true")
    ap.add_argument("--also-docx", action="store_true",
                    help="also emit a readable .docx view (outputs/docx/)")
    ap.add_argument("--include-body", action="store_true",
                    help="include full prose body text in JSON (default: metadata only)")
    ap.add_argument("--all-formats", action="store_true",
                    help="shortcut: equivalent to --also-markdown --also-docx --include-body")
    args = ap.parse_args()

    if args.all_formats:
        args.also_markdown = True
        args.also_docx = True
        args.include_body = True

    conn = open_db(args.db)
    extract = build_extract(conn, include_body=args.include_body)

    out_dir = Path(OUT_DIR)
    out_dir.mkdir(parents=True, exist_ok=True)
    stamp = today_compact()
    out_path = Path(args.out) if args.out else out_dir / f"wa-programme-prose-extract-{stamp}.json"
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(json.dumps(extract, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Wrote JSON: {out_path}")

    if args.also_markdown:
        md_path = out_path.with_suffix(".md")
        md_path.write_text(render_markdown_view(extract), encoding="utf-8")
        print(f"Wrote MD:   {md_path}")

    if args.also_docx:
        docx_out_dir = Path(DOCX_OUT_DIR)
        docx_out_dir.mkdir(parents=True, exist_ok=True)
        docx_path = docx_out_dir / f"wa-programme-prose-extract-{stamp}.docx"
        if not args.include_body:
            print("[NOTE] --also-docx implies --include-body; rebuilding extract with bodies…")
            extract_with_bodies = build_extract(conn, include_body=True)
            render_docx(extract_with_bodies, docx_path)
        else:
            render_docx(extract, docx_path)
        print(f"Wrote DOCX: {docx_path}")

    pp = extract["programme_prose"]
    print(f"Section types: {pp['type_count']}  Content sections: {pp['section_total']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
