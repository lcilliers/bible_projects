"""Render a single integrated-essay markdown file to .docx.

Companion to combine_cluster_published_to_docx.py (which combines chapter drafts).
This renders a single essay-form md (the v3_0 Phase E publication target,
per §11.3 step 5) into the publication .docx.

Usage:
    python scripts/_render_essay_to_docx_v1_20260527.py \
        --in Sessions/Session_Clusters/M11/publishing/wa-cluster-M11-essay-v1-20260527.md

Output: .docx sibling to the .md input.

Handles: headings (#, ##, ###), italics (*..*), bold (**..**), horizontal rules,
blockquotes, plain paragraphs. Strips administrative front-matter block (the
"**File:**" + "**Programme reference:**" + "**Authored:**" lines).
"""
from __future__ import annotations
import argparse, io, re, sys
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches

RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")
RE_HR = re.compile(r"^---+\s*$")
RE_HEADING = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
RE_FRONT_MATTER = re.compile(r"^\*\*(File|Programme reference|Authored|Cluster):\*\*", re.IGNORECASE)


def add_inline_runs(para, text: str) -> None:
    """Process **bold** and *italic* into runs."""
    cursor = 0
    while cursor < len(text):
        next_bold = RE_BOLD.search(text, cursor)
        next_italic = RE_ITALIC.search(text, cursor)
        candidates = [m for m in (next_bold, next_italic) if m is not None]
        if not candidates:
            para.add_run(text[cursor:])
            break
        m = min(candidates, key=lambda x: x.start())
        if m.start() > cursor:
            para.add_run(text[cursor:m.start()])
        run = para.add_run(m.group(1))
        if m.re is RE_BOLD:
            run.bold = True
        else:
            run.italic = True
        cursor = m.end()


def render(in_path: Path, out_path: Path) -> None:
    text = in_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    # Body style sizing
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Strip admin front-matter lines
        if RE_FRONT_MATTER.match(stripped):
            i += 1
            continue

        # Horizontal rule -> page break or visual gap
        if RE_HR.match(stripped):
            doc.add_paragraph()
            i += 1
            continue

        # Heading
        m = RE_HEADING.match(stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            heading_text = RE_BOLD.sub(r"\1", heading_text)
            heading_text = RE_ITALIC.sub(r"\1", heading_text)
            doc.add_heading(heading_text, level=min(level, 4))
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                quote_lines.append(re.sub(r"^\s*>\s?", "", lines[i]).rstrip())
                i += 1
            quote_text = " ".join(q for q in quote_lines if q.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            add_inline_runs(p, quote_text)
            for r in p.runs:
                r.italic = True
            continue

        # Normal paragraph (accumulate consecutive non-blank non-special lines)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if not ns:
                break
            if RE_HEADING.match(ns) or RE_HR.match(ns) or ns.startswith(">"):
                break
            if RE_FRONT_MATTER.match(ns):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_inline_runs(p, para_text)

    doc.save(out_path)
    print(f"Wrote {out_path}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_path", required=True, type=Path)
    ap.add_argument("--out", dest="out_path", default=None, type=Path)
    args = ap.parse_args()

    in_path = args.in_path
    if not in_path.exists():
        print(f"FAIL: input not found: {in_path}")
        return 1
    out_path = args.out_path or in_path.with_suffix(".docx")
    render(in_path, out_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
