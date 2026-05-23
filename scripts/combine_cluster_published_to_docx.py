"""Combine the latest chapter drafts in a cluster's Published/ folder into one .docx.

Usage:
  python scripts/combine_cluster_published_to_docx.py --cluster M01
  python scripts/combine_cluster_published_to_docx.py --cluster M01 --chapters 1-7
  python scripts/combine_cluster_published_to_docx.py --cluster M01 --include-appendices

Reads `Sessions/Session_Clusters/{CLUSTER}/Published/wa-cluster-{CLUSTER}-ch{N}-draft-v{V}-{date}.md`
files, picks the highest v per chapter (and the latest date if multiple v's share a number),
and writes a single combined .docx into the same Published/ folder with canonical name:

  wa-cluster-{CLUSTER}-{short_name}-combined-v{N}-{YYYYMMDD}.docx

Cluster description and short_name are sourced from the `cluster` row in the DB.

Per-chapter preprocessing (so the combined doc is publication-ready, not analytical):
  - The chapter's H1 generic title (e.g. `# Fear, Dread and Terror — Chapter 1 input`)
    is dropped; the H2 line (e.g. `## What this study is`) is used as the **topical
    chapter title** in the combined doc (rendered as Heading 1).
  - The metadata block (`**Cluster:** … **Chapter:** … **Generated:** …`) is dropped.
  - The "Cross-chapter consistency: characteristics in this study" block (where present)
    is dropped — it is reference material for the AI session, not for the published doc.
  - The redundant numbered chapter header (`## 1. What this study is`) is dropped because
    the topical title already serves as the chapter heading.
  - `<!-- EVIDENCE: ... -->` … `<!-- /EVIDENCE -->` blocks are dropped — they are the
    per-section verse-evidence packets included for AI grounding, not for the published doc.

A Table of Contents listing the topical title of every chapter is inserted on the first
page (static; survives editing without needing Word to refresh fields).

Other markdown elements handled:
  - headings (#, ##, ###, ####)
  - paragraphs with inline **bold** and *italic* / _italic_
  - bullet lists (- item)
  - blockquotes (> ...)
  - horizontal rules (---) -> visual paragraph break
  - links [text](url) -> plain text (URLs dropped)
  - other HTML comments stripped
A hard page break is inserted between chapters and appendices.
"""
from __future__ import annotations

import argparse
import re
import sqlite3
import sys
import io
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

REPO = Path(__file__).resolve().parent.parent
DB = REPO / "database" / "bible_research.db"

# Inline markdown patterns
RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC_AST = re.compile(r"(?<!\*)\*([^*\n]+?)\*(?!\*)")
RE_ITALIC_UND = re.compile(r"(?<![A-Za-z0-9_])_([^_\n]+?)_(?![A-Za-z0-9_])")
RE_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
RE_HTML_COMMENT = re.compile(r"<!--.*?-->", re.DOTALL)

# Per-chapter preprocessing patterns
RE_HR = re.compile(r"^---+\s*$")
RE_EVIDENCE_OPEN = re.compile(r"<!--\s*EVIDENCE[\s:]", re.IGNORECASE)
RE_EVIDENCE_CLOSE = re.compile(r"<!--\s*/\s*EVIDENCE", re.IGNORECASE)
RE_CROSS_CHAPTER_HEADER = re.compile(
    r"^##\s+Cross-chapter\s+consistency", re.IGNORECASE
)
RE_NUMBERED_SECTION_HEADER = re.compile(r"^##\s+\d+\.\s+")
# Bold-line analytical labels that introduce evidence content but live outside the
# <!-- EVIDENCE --> comment delimiters. Drop the whole line.
RE_ANALYTICAL_LABEL = re.compile(
    r"^\s*\*\*\s*Evidence\b[^*]*\*\*\s*$", re.IGNORECASE
)

# Chapter / appendix filename patterns
RE_CHAPTER_FILE = re.compile(
    r"^wa-cluster-([A-Za-z0-9]+)-ch(\d+)-draft-v(\d+)-(\d{8})\.md$", re.IGNORECASE
)
RE_APPENDIX_FILE = re.compile(
    r"^wa-cluster-([A-Za-z0-9]+)-app([a-z])-draft-v(\d+)-(\d{8})\.md$", re.IGNORECASE
)


def parse_chapter_range(spec: str) -> list[int]:
    """Parse '1-7' or '1,3,5' or '1,3-5' into a sorted list of chapter numbers."""
    out: set[int] = set()
    for part in spec.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            out.update(range(int(a), int(b) + 1))
        else:
            out.add(int(part))
    return sorted(out)


def discover_latest_drafts(
    src_dir: Path,
    cluster_code: str,
    chapters: list[int] | None,
    include_appendices: bool,
) -> list[tuple[str, Path]]:
    """Return ordered list of (label, path) for the chapters / appendices to combine.

    Picks the highest version per chapter; among same-version files picks the latest date.
    """
    by_chapter: dict[int, list[tuple[int, str, Path]]] = defaultdict(list)
    by_appendix: dict[str, list[tuple[int, str, Path]]] = defaultdict(list)

    for p in src_dir.iterdir():
        if not p.is_file():
            continue
        mc = RE_CHAPTER_FILE.match(p.name)
        if mc and mc.group(1).upper() == cluster_code.upper():
            n = int(mc.group(2))
            v = int(mc.group(3))
            d = mc.group(4)
            by_chapter[n].append((v, d, p))
            continue
        ma = RE_APPENDIX_FILE.match(p.name)
        if ma and ma.group(1).upper() == cluster_code.upper():
            letter = ma.group(2).lower()
            v = int(ma.group(3))
            d = ma.group(4)
            by_appendix[letter].append((v, d, p))

    if not by_chapter:
        raise SystemExit(f"No chapter files found in {src_dir} for cluster {cluster_code}.")

    selected_chapters = sorted(by_chapter.keys()) if chapters is None else chapters
    out: list[tuple[str, Path]] = []
    for n in selected_chapters:
        if n not in by_chapter:
            raise SystemExit(f"Chapter {n} not found for {cluster_code} in {src_dir}.")
        v, d, p = max(by_chapter[n], key=lambda t: (t[0], t[1]))
        out.append((f"Chapter {n} (v{v}, {d})", p))

    if include_appendices and by_appendix:
        for letter in sorted(by_appendix.keys()):
            v, d, p = max(by_appendix[letter], key=lambda t: (t[0], t[1]))
            out.append((f"Appendix {letter.upper()} (v{v}, {d})", p))

    return out


def next_output_version(dst_dir: Path, base_name: str, date_str: str) -> int:
    pat = re.compile(rf"^{re.escape(base_name)}-v(\d+)-{date_str}\.docx$", re.IGNORECASE)
    max_v = 0
    if dst_dir.exists():
        for p in dst_dir.iterdir():
            m = pat.match(p.name)
            if m:
                max_v = max(max_v, int(m.group(1)))
    return max_v + 1


def add_runs_with_inline_formatting(paragraph, text: str) -> None:
    """Tokenise text into bold / italic / link / plain segments and add as runs."""
    text = RE_LINK.sub(r"\1", text)

    BOLD_OPEN, BOLD_CLOSE = "\x01B<", "\x01b>"
    ITAL_OPEN, ITAL_CLOSE = "\x02I<", "\x02i>"

    text = RE_BOLD.sub(lambda m: f"{BOLD_OPEN}{m.group(1)}{BOLD_CLOSE}", text)
    text = RE_ITALIC_AST.sub(lambda m: f"{ITAL_OPEN}{m.group(1)}{ITAL_CLOSE}", text)
    text = RE_ITALIC_UND.sub(lambda m: f"{ITAL_OPEN}{m.group(1)}{ITAL_CLOSE}", text)

    token_re = re.compile(r"(\x01B<|\x01b>|\x02I<|\x02i>)")
    parts = token_re.split(text)
    bold = False
    italic = False
    for p in parts:
        if p == BOLD_OPEN:
            bold = True
        elif p == BOLD_CLOSE:
            bold = False
        elif p == ITAL_OPEN:
            italic = True
        elif p == ITAL_CLOSE:
            italic = False
        elif p:
            run = paragraph.add_run(p)
            run.bold = bold
            run.italic = italic


def preprocess_chapter(md: str) -> tuple[str, str]:
    """Strip chapter preamble, cross-chapter-consistency block, redundant numbered
    section header, and evidence blocks. Return (topical_title, cleaned_body_md).

    The topical title comes from the H2 line that immediately follows the H1 line.
    Body content starts after the first `---` separator that follows the metadata.
    If the next H2 after that separator is the cross-chapter-consistency block, it
    is dropped together with its trailing `---`. The redundant `## N. <title>`
    header is then also dropped (the topical title already serves as the heading).
    `<!-- EVIDENCE: ... --> ... <!-- /EVIDENCE -->` blocks are excised line-by-line.
    """
    lines = md.splitlines()

    # 1. Find H2 topical title (the H2 line nearest the top, after the H1)
    h2_idx = None
    for i, line in enumerate(lines):
        s = line.strip()
        if s.startswith("## ") and not s.startswith("## #"):
            h2_idx = i
            break
    if h2_idx is None:
        raise ValueError("No H2 topical title found in chapter")
    topical_title = lines[h2_idx][3:].strip()
    for pat in (RE_BOLD, RE_ITALIC_AST, RE_ITALIC_UND):
        topical_title = pat.sub(r"\1", topical_title)
    topical_title = RE_LINK.sub(r"\1", topical_title)

    # 2. Find first `---` after H2 — marks end of metadata block
    sep_idx = None
    for i in range(h2_idx + 1, len(lines)):
        if RE_HR.match(lines[i].strip()):
            sep_idx = i
            break
    if sep_idx is None:
        raise ValueError("No '---' separator found after chapter metadata")
    body_start = sep_idx + 1

    # 3. If next non-empty line is cross-chapter-consistency H2, skip its whole block
    j = body_start
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and RE_CROSS_CHAPTER_HEADER.match(lines[j].strip()):
        # Skip to next `---`
        for k in range(j + 1, len(lines)):
            if RE_HR.match(lines[k].strip()):
                body_start = k + 1
                break

    # 4. If the next non-empty line is `## N. <title>`, skip it (redundant header)
    j = body_start
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and RE_NUMBERED_SECTION_HEADER.match(lines[j].strip()):
        body_start = j + 1
        # Skip any trailing blanks immediately after
        while body_start < len(lines) and not lines[body_start].strip():
            body_start += 1

    # 5. Drop evidence blocks (HTML-comment-delimited, possibly multi-line) and
    #    drop bold-line analytical labels like `**Evidence: cluster-wide claims …**`.
    out_lines: list[str] = []
    in_evidence = False
    for line in lines[body_start:]:
        if not in_evidence and RE_EVIDENCE_OPEN.search(line):
            in_evidence = True
            continue
        if in_evidence:
            if RE_EVIDENCE_CLOSE.search(line):
                in_evidence = False
            continue
        if RE_ANALYTICAL_LABEL.match(line):
            continue
        out_lines.append(line)

    return topical_title, "\n".join(out_lines)


def add_toc(doc: Document, sections: list[tuple[int, str, str]]) -> None:
    """Insert a simple, static TOC. `sections` is a list of (chapter_num, topical_title, kind)
    where kind is 'chapter' or 'appendix'. Page numbers are deliberately omitted (would
    require post-render computation) — the TOC names titles and lets the reader navigate
    via the chapter headings.
    """
    doc.add_heading("Table of Contents", level=1)
    for num, title, kind in sections:
        p = doc.add_paragraph()
        if kind == "appendix":
            label_run = p.add_run(f"Appendix {num}.  ")
        else:
            label_run = p.add_run(f"Chapter {num}.  ")
        label_run.bold = True
        p.add_run(title)


def render_markdown_to_docx(doc: Document, md: str) -> None:
    md = RE_HTML_COMMENT.sub("", md)
    lines = md.splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r"^---+\s*$", stripped):
            doc.add_paragraph()
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.+?)\s*#*\s*$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            for pat in (RE_BOLD, RE_ITALIC_AST, RE_ITALIC_UND):
                text = pat.sub(r"\1", text)
            text = RE_LINK.sub(r"\1", text)
            doc.add_heading(text, level=min(level, 4))
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                bullet_text = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_inline_formatting(p, bullet_text)
                i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].lstrip().startswith(">"):
                ql = re.sub(r"^\s*>\s?", "", lines[i])
                quote_lines.append(ql.rstrip())
                i += 1
            quote_text = " ".join(q for q in quote_lines if q.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.right_indent = Inches(0.4)
            add_runs_with_inline_formatting(p, quote_text)
            for r in p.runs:
                r.italic = True
            continue

        # Normal paragraph
        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            ns = nxt.strip()
            if not ns:
                break
            if re.match(r"^(#{1,6}\s|>|[-*]\s|---+\s*$)", ns):
                break
            para_lines.append(nxt)
            i += 1
        para_text = " ".join(l.strip() for l in para_lines)
        p = doc.add_paragraph()
        add_runs_with_inline_formatting(p, para_text)


def fetch_cluster_meta(cluster_code: str) -> tuple[str, str]:
    conn = sqlite3.connect(f"file:{DB}?mode=ro", uri=True, timeout=5.0)
    try:
        r = conn.execute(
            "SELECT short_name, description FROM cluster WHERE cluster_code=?",
            (cluster_code,),
        ).fetchone()
    finally:
        conn.close()
    if not r:
        raise SystemExit(f"cluster {cluster_code} not found in DB")
    return r[0], r[1]


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    ap.add_argument("--cluster", required=True,
                    help="Cluster code, e.g. M01")
    ap.add_argument("--chapters", default=None,
                    help="Chapter range to include, e.g. '1-7' or '1,3-5'. "
                         "Default: all chapter files found.")
    ap.add_argument("--include-appendices", action="store_true",
                    help="Append discovered appendix files (wa-cluster-{CODE}-app{a-z}-*) "
                         "after the chapters.")
    ap.add_argument("--out", type=Path, default=None,
                    help="Explicit output path. Default: "
                         "Sessions/Session_Clusters/{CLUSTER}/Published/"
                         "wa-cluster-{CLUSTER}-{shortname}-combined-v{N}-{YYYYMMDD}.docx")
    args = ap.parse_args()

    cluster_code = args.cluster.strip()
    short_name, description = fetch_cluster_meta(cluster_code)

    src_dir = REPO / "Sessions" / "Session_Clusters" / cluster_code / "Published"
    if not src_dir.exists():
        raise SystemExit(f"Published folder not found: {src_dir}")

    chapters = parse_chapter_range(args.chapters) if args.chapters else None
    pieces = discover_latest_drafts(src_dir, cluster_code, chapters, args.include_appendices)

    date_str = datetime.now(timezone.utc).strftime("%Y%m%d")
    short_slug = re.sub(r"[^A-Za-z0-9]+", "-", short_name.strip()).strip("-") or "cluster"
    base_name = f"wa-cluster-{cluster_code}-{short_slug}-combined"
    if args.out:
        out_path = args.out
    else:
        v = next_output_version(src_dir, base_name, date_str)
        out_path = src_dir / f"{base_name}-v{v}-{date_str}.docx"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # First pass: preprocess each chapter to extract topical title + cleaned body
    prepared: list[tuple[str, str, Path, str, str]] = []  # (label, kind, path, topical_title, body_md)
    for label, path in pieces:
        md = path.read_text(encoding="utf-8")
        title, body = preprocess_chapter(md)
        kind = "appendix" if label.startswith("Appendix") else "chapter"
        prepared.append((label, kind, path, title, body))

    # Build docx
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title page
    doc.add_heading(description, level=0)
    subtitle = doc.add_paragraph()
    sr = subtitle.add_run(
        f"A study of the Bible's inner-life vocabulary — "
        f"{cluster_code} cluster · {len(prepared)} sections · "
        f"compiled {date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"
    )
    sr.italic = True
    doc.add_paragraph()
    doc.add_paragraph(
        f"Combined draft compiled from "
        f"Sessions/Session_Clusters/{cluster_code}/Published/."
    )
    doc.add_paragraph()

    # TOC on the first page
    chapter_counter = 0
    appendix_counter = 0
    toc_entries: list[tuple[str | int, str, str]] = []
    for label, kind, _path, title, _body in prepared:
        if kind == "chapter":
            chapter_counter += 1
            toc_entries.append((chapter_counter, title, "chapter"))
        else:
            appendix_counter += 1
            letter = chr(ord("A") + appendix_counter - 1)
            toc_entries.append((letter, title, "appendix"))
    add_toc(doc, toc_entries)

    # Page break before first section
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

    chapter_counter = 0
    appendix_counter = 0
    for idx, (label, kind, path, topical_title, body_md) in enumerate(prepared, start=1):
        print(f"{label}: {path.name} ({len(body_md):,} chars after preprocessing)")
        if kind == "chapter":
            chapter_counter += 1
            heading_text = f"Chapter {chapter_counter}.  {topical_title}"
        else:
            appendix_counter += 1
            letter = chr(ord("A") + appendix_counter - 1)
            heading_text = f"Appendix {letter}.  {topical_title}"
        doc.add_heading(heading_text, level=1)
        render_markdown_to_docx(doc, body_md)
        if idx < len(prepared):
            p = doc.add_paragraph()
            p.add_run().add_break(WD_BREAK.PAGE)

    doc.save(out_path)
    print(f"\nWrote: {out_path.relative_to(REPO)}  ({out_path.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
