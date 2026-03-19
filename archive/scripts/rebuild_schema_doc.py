"""
rebuild_schema_doc.py
─────────────────────
Regenerates docs/DB-Schema-Overview.docx from the live database.

Run:
    python scripts/rebuild_schema_doc.py
"""
import datetime
import os
import sys

ROOT_DIR = os.path.join(os.path.dirname(__file__), "..")
sys.path.insert(0, ROOT_DIR)

from analytics.db_client import get_connection  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402

# ── Helpers ────────────────────────────────────────────────────────────────────

def _set_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(tbl, headers, header_color="1F497D", col_widths=None):
    row = tbl.rows[0]
    for j, h in enumerate(headers):
        c = row.cells[j]
        if col_widths:
            c.width = Cm(col_widths[j])
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_bg(c, header_color)


def _data_table(doc, headers, rows, header_color="1F497D", col_widths=None):
    if not rows:
        doc.add_paragraph("(no data)").runs[0].font.size = Pt(9)
        return
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    _header_row(tbl, headers, header_color, col_widths)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i + 1].cells[j]
            if col_widths:
                c.width = Cm(col_widths[j])
            r = c.paragraphs[0].add_run("" if val is None else str(val))
            r.font.size = Pt(8)
            if i % 2 == 0:
                _set_bg(c, "EEF3F8")


def _kv_table(doc, data, col_widths=(5, 11.5)):
    tbl = doc.add_table(rows=len(data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(data):
        kc = tbl.rows[i].cells[0]
        kc.width = Cm(col_widths[0])
        kr = kc.paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(8)
        _set_bg(kc, "DCE6F1")
        vc = tbl.rows[i].cells[1]
        vc.width = Cm(col_widths[1])
        vr = vc.paragraphs[0].add_run("" if v is None else str(v))
        vr.font.size = Pt(8)


def _note(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def _divider(doc):
    doc.add_paragraph("─" * 90).runs[0].font.size = Pt(7)


# ── Schema queries ─────────────────────────────────────────────────────────────

def load_schema(conn):
    """Return all schema facts from the live DB."""
    tables_raw = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' "
        "AND name != 'sqlite_sequence' ORDER BY name"
    ).fetchall()
    tables = [r[0] for r in tables_raw]

    schema = {}
    for t in tables:
        cols = [dict(r) for r in conn.execute(f"PRAGMA table_info({t})").fetchall()]
        fks  = [dict(r) for r in conn.execute(f"PRAGMA foreign_key_list({t})").fetchall()]
        idxs_raw = conn.execute(f"PRAGMA index_list({t})").fetchall()
        idxs = []
        for idx in idxs_raw:
            sql_row = conn.execute(
                "SELECT sql FROM sqlite_master WHERE name=?", (idx["name"],)
            ).fetchone()
            idxs.append({
                "name": idx["name"],
                "unique": idx["unique"],
                "sql": sql_row[0] if sql_row else None,
            })
        count = conn.execute(f"SELECT COUNT(*) FROM {t}").fetchone()[0]
        schema[t] = dict(cols=cols, fks=fks, idxs=idxs, count=count)
    return schema


# ── Document sections ──────────────────────────────────────────────────────────

# Display order for tables — grouped by subsystem
TABLE_GROUPS = [
    ("Word Registry",               ["word_registry"]),
    ("MTI — Master Term Index",     ["mti_terms", "mti_term_flags", "mti_term_cross_refs"]),
    ("WA — Word Analysis",          ["wa_file_index", "wa_term_inventory",
                                     "wa_term_related_words", "wa_term_root_family",
                                     "wa_term_phase2_flags", "phase2_flag_types",
                                     "wa_cross_registry_links", "wa_crosslink_type",
                                     "wa_data_quality_flags",
                                     "wa_quality_flag_types", "wa_verse_records"]),
    ("Reference / Future Use",      ["books", "sources", "themes"]),
]

HEADER_COLORS = {
    "Word Registry":             "2F5496",
    "MTI — Master Term Index":   "375623",
    "WA — Word Analysis":        "833C00",
    "Reference / Future Use":    "595959",
}


def build_doc(schema: dict) -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # ── Title ──────────────────────────────────────────────────────────────
    tp = doc.add_heading("Bible Research Database — Schema Overview", 0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
        f"  |  Database: bible_research.db  |  SQLite"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # S1 — TABLE RECORD COUNTS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Table Record Counts", 1)

    summary_rows = []
    for group_name, group_tables in TABLE_GROUPS:
        for t in group_tables:
            if t in schema:
                n = schema[t]["count"]
                status = ("Populated" if n > 0 else "Empty — reserved for future use")
                summary_rows.append([group_name, t, f"{n:,}", status])

    _data_table(doc,
        ["Group", "Table", "Rows", "Status"],
        summary_rows,
        header_color="2F5496",
        col_widths=[4.5, 5, 2.5, 7])
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # S2 — ENTITY RELATIONSHIPS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2. Entity Relationships", 1)

    wr_count   = schema["word_registry"]["count"]
    wfi_count  = schema["wa_file_index"]["count"]
    wti_count  = schema["wa_term_inventory"]["count"]
    rw_count   = schema["wa_term_related_words"]["count"]
    rf_count   = schema["wa_term_root_family"]["count"]
    p2f_count  = schema["wa_term_phase2_flags"]["count"]
    vr_count   = schema["wa_verse_records"]["count"]
    mti_count  = schema["mti_terms"]["count"]
    xcr_count  = schema["mti_term_cross_refs"]["count"]
    mf_count   = schema["mti_term_flags"]["count"]
    p2ft_count = schema["phase2_flag_types"]["count"]
    qft_count  = schema["wa_quality_flag_types"]["count"]
    crl_count  = schema["wa_cross_registry_links"]["count"]
    clt_count  = schema["wa_crosslink_type"]["count"]
    dqf_count  = schema["wa_data_quality_flags"]["count"]
    bk_count   = schema["books"]["count"]

    erd_text = (
        f"word_registry ({wr_count:,} rows)\n"
        f"   no   referenced by mti_terms.owning_registry (string match)\n"
        f"   id   wa_file_index.word_registry_fk\n\n"

        f"wa_file_index ({wfi_count:,} rows)   [one row per source JSON file / part]\n"
        f"   id   wa_term_inventory.file_id\n"
        f"            wa_verse_records.file_id\n"
        f"            wa_cross_registry_links.file_id\n"
        f"            wa_data_quality_flags.file_id\n\n"
        f"wa_cross_registry_links ({crl_count:,} rows)\n"
        f"   connection_type_id    wa_crosslink_type.id   ({clt_count:,} type codes)\n"
        f"   linked_registry_id   word_registry.id   (nullable; 7 unresolvable stay NULL)\n"
        f"   id   mti_terms.word_data_reference  "
        f"(522 linked; 107 NULL = archived words)\n\n"

        f"wa_term_inventory ({wti_count:,} rows)   [one row per Hebrew/Greek term per file]\n"
        f"   id   wa_term_related_words.term_inv_id   ({rw_count:,} rows)\n"
        f"   id   wa_term_root_family.term_inv_id     ({rf_count:,} rows)\n"
        f"   id   wa_term_phase2_flags.term_inv_id    ({p2f_count:,} rows, junction)\n"
        f"   id   wa_verse_records.term_inv_id    ({vr_count:,} rows)\n\n"

        f"wa_term_phase2_flags  [many-to-many junction]\n"
        f"   flag_id    phase2_flag_types.id   ({p2ft_count:,} rows, reference)\n\n"

        f"wa_data_quality_flags\n"
        f"   flag_id    wa_quality_flag_types.id   ({qft_count:,} rows, reference)\n\n"

        f"mti_terms ({mti_count:,} rows)   [Master Term Index]\n"
        f"   id   mti_term_cross_refs.mti_term_id     ({xcr_count:,} rows)\n"
        f"   id   mti_term_flags.mti_term_id          ({mf_count:,} rows, junction)\n"
        f"   word_data_reference  wa_file_index.id\n\n"

        f"mti_term_flags  [many-to-many junction]\n"
        f"   flag_id    phase2_flag_types.id   ({p2ft_count:,} rows, reference)\n\n"

        f"wa_verse_records ({vr_count:,} rows)\n"
        f"   book_id    books.id   ({bk_count:,} rows — 66 canonical Bible books)\n\n"

        f"Reference tables (currently empty — reserved for future annotation)\n"
        f"   books ({bk_count:,} rows — populated with 66 canonical books)\n"
        f"   sources ({schema['sources']['count']:,} rows)   themes ({schema['themes']['count']:,} rows)"
    )

    p = doc.add_paragraph(erd_text)
    p.runs[0].font.name = "Courier New"
    p.runs[0].font.size = Pt(8.5)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # S3 — TABLE SCHEMAS (per group)
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3. Table Schemas", 1)

    for group_name, group_tables in TABLE_GROUPS:
        group_color = HEADER_COLORS[group_name]
        doc.add_heading(f"3.x  {group_name}", 2).runs[0].font.color.rgb = RGBColor(
            int(group_color[0:2], 16),
            int(group_color[2:4], 16),
            int(group_color[4:6], 16),
        )

        for tname in group_tables:
            if tname not in schema:
                continue
            s = schema[tname]
            doc.add_heading(
                f"{tname}   ({s['count']:,} rows)", 3
            )

            # Columns table
            col_rows = []
            for c in s["cols"]:
                pk  = "PK"  if c["pk"]     else ""
                nn  = "✓"   if c["notnull"] else ""
                df  = c["dflt_value"] or ""
                col_rows.append([c["cid"], c["name"], c["type"], nn, pk, df])

            _data_table(doc,
                ["#", "Column", "Type", "Not Null", "PK", "Default"],
                col_rows,
                header_color=group_color,
                col_widths=[0.8, 5.5, 2.5, 1.8, 1, 3.5])

            # Foreign keys
            if s["fks"]:
                fk_str = "  |  ".join(
                    f"[{fk['from']}]  →  {fk['table']}.{fk['to']}"
                    for fk in s["fks"]
                )
                doc.add_paragraph(f"Foreign Keys:  {fk_str}"
                                   ).runs[0].font.size = Pt(8)

            # Indexes (skip auto-generated ones)
            named_idxs = [
                i for i in s["idxs"]
                if i["sql"] is not None  # skip sqlite_autoindex_*
            ]
            if named_idxs:
                idx_lines = []
                for idx in named_idxs:
                    uniq = "  [UNIQUE]" if idx["unique"] else ""
                    # Extract just the ON clause for compactness
                    sql_clean = (idx["sql"] or "").replace("\n", " ").replace("    ", " ")
                    idx_lines.append(f"  {idx['name']}{uniq}\n      {sql_clean}")
                doc.add_paragraph("Indexes:\n" + "\n".join(idx_lines)
                                   ).runs[0].font.size = Pt(8)

            # Special notes for wa_verse_records trigger
            if tname == "wa_verse_records":
                _note(doc,
                    "Trigger:  wa_verse_records_updated_at — keeps updated_at current "
                    "on every UPDATE via AFTER UPDATE trigger.")

            doc.add_paragraph()

        _divider(doc)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # S4 — DESIGN NOTES
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4. Design Notes", 1)

    design_notes = [
        ("Column / Area", "Note"),
        ("mti_terms.strongs_number",
         "Nullable and non-unique. Two known legitimate duplicates: G0150 (shame — "
         "two inflected forms) and H7999A (peace — consolidated sub-gloss)."),
        ("mti_terms.word_data_reference",
         "Links to wa_file_index.id. 522 of 629 terms linked. "
         "107 NULL = distress, joy, peace files archived for rebuild."),
        ("mti_terms.owning_part",
         "Stored as raw TEXT (mixed formats from source JSON: null, \"1\", \"Part1\", integer). "
         "word_data_reference is the clean resolved FK."),
        ("wa_file_index.word_registry_fk",
         "Links to word_registry.id. All 41 files fully resolved (0 unmatched)."),
        ("wa_verse_records.term_inv_id",
         "Links to wa_term_inventory.id. 1 NULL row — abomination file contained a "
         "term with no reference/book."),
        ("wa_verse_records.book_id",
         "FK to books.id, populated during the books-migration pass (all 66 books loaded). "
         "1 bad-data row has NULL book_id (same row as NULL term_inv_id above)."),
        ("wa_verse_records.translation",
         "NOT NULL DEFAULT 'ESV'. Populated from wa_file_index.translation_used at import."),
        ("wa_verse_records columns: note / claude_output",
         "Research annotation fields — note for researcher text, claude_output for the raw "
         "Claude JSON response. Both NULL until Phase 2 annotation begins."),
        ("wa_verse_records: updated_at trigger",
         "AFTER UPDATE trigger wa_verse_records_updated_at keeps updated_at in sync. "
         "created_at is set at import time; updated_at reflects last annotation change."),
        ("idx_wavr_file_term_pos",
         "Composite index on wa_verse_records (file_id, term_id, book_id, chapter, verse_num). "
         "Added 2026-03-16. Eliminates temp B-tree sort for single-file-id extract queries. "
         "For multi-part words (IN list), Python-side sort is used instead."),
        ("phase2_flag_types  (consolidated flag lookup)",
         "Single lookup for all phase-2 analytical signals (14 codes). "
         "Replaces the former wa_phase2_flag_types (12 codes) and mti_phase2_flags (1 code). "
         "Shared by wa_term_phase2_flags and mti_term_flags via flag_id FK. "
         "Phase-2 flags signal terms that require deeper analysis in the meaning-derivation phase."),
        ("wa_quality_flag_types  (quality flag lookup)",
         "Two-level quality flag lookup: flag_group / flag_code (22 codes across 4 groups: "
         "DATA_COVERAGE, IMPORT_STATUS, TERM_ANALYSIS, NOTE). "
         "Multiple flag_code values can share the same flag_group (e.g. different kinds of missing data)."),
        ("wa_data_quality_flags  (normalised 2026-03-16)",
         "Free-text flag column replaced by flag_id INTEGER FK → wa_quality_flag_types. "
         "Slash-combined entries split into separate rows. "
         "PERIPHERAL_TERM / SOMATIC_EXPRESSION: quality row kept as PERIPHERAL_TERM; "
         "SOMATIC_EXPRESSION added to wa_term_phase2_flags for the relevant term."),
        ("phase2_flags normalisation",
         "All term-level phase2_flags normalised to UPPER_SNAKE_CASE at import time."),
        ("wa_term_inventory inline boolean flags",
         "god_as_subject, somatic_link, causative_form_present are stored as INTEGER "
         "booleans (0/1) directly on wa_term_inventory in addition to the normalised "
         "wa_term_phase2_flags junction table. Both are kept in sync at import."),
        ("books table",
         "Now fully populated: 66 rows covering all canonical Bible books. "
         "Referenced by wa_verse_records.book_id."),
        ("sources / themes",
         "Original schema scaffolding, currently empty. Reserved for future Phase 2 "
         "annotation and cross-reference work."),
        ("Archived files",
         "Distress (parts 1–3), Joy (parts 1–3), Peace (parts 1–3) moved to Archive "
         "folder. These 9 files will be rebuilt using the current schema version."),
        ("Superseded duplicates",
         "Three non-canonical files kept in Session_A_Data folder for reference "
         "(anguish original, grief part1/part2 pre-update). Import used updated versions only."),
        ("wa_verse_records index duplication",
         "idx_wa_vr_ref and idx_wavr_reference both index the reference column (legacy + schema). "
         "Similarly idx_wa_vr_book and idx_wavr_book_ch_v overlap. "
         "Consolidation deferred — no functional impact."),
    ]

    # Header row first
    tbl = doc.add_table(rows=len(design_notes), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, (k, v) in enumerate(design_notes):
        kc = tbl.rows[i].cells[0]
        kc.width = Cm(5.5)
        vc = tbl.rows[i].cells[1]
        vc.width = Cm(11)
        if i == 0:
            kr = kc.paragraphs[0].add_run(k)
            kr.bold = True
            kr.font.size = Pt(8)
            kr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_bg(kc, "1F497D")

            vr = vc.paragraphs[0].add_run(v)
            vr.bold = True
            vr.font.size = Pt(8)
            vr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_bg(vc, "1F497D")
        else:
            kr = kc.paragraphs[0].add_run(k)
            kr.bold = True
            kr.font.size = Pt(8)
            _set_bg(kc, "DCE6F1")

            vr = vc.paragraphs[0].add_run(v)
            vr.font.size = Pt(8)
            if i % 2 == 0:
                _set_bg(vc, "EEF3F8")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # S5 — INDEX SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5. Full Index Inventory", 1)
    _note(doc, "All named indexes across all tables. Auto-generated UNIQUE indexes "
               "(sqlite_autoindex_*) are omitted.")
    doc.add_paragraph()

    idx_summary_rows = []
    for group_name, group_tables in TABLE_GROUPS:
        for tname in group_tables:
            if tname not in schema:
                continue
            for idx in schema[tname]["idxs"]:
                if idx["sql"] is None:
                    continue  # skip auto
                uniq = "Yes" if idx["unique"] else ""
                sql_clean = (idx["sql"] or "").replace("\n", " ").strip()
                idx_summary_rows.append([tname, idx["name"], uniq, sql_clean])

    _data_table(doc,
        ["Table", "Index Name", "Unique", "Definition"],
        idx_summary_rows,
        header_color="2F5496",
        col_widths=[4.5, 5, 1.5, 10.5])

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    conn = get_connection()
    print("Loading schema from live DB …")
    schema = load_schema(conn)
    conn.close()

    table_list = sorted(schema.keys())
    print(f"  Tables: {table_list}")
    for t, s in sorted(schema.items()):
        print(f"  {t}: {s['count']:,} rows, {len(s['cols'])} cols, "
              f"{len([i for i in s['idxs'] if i['sql']])} named indexes")

    print("Building document …")
    doc = build_doc(schema)

    out_path = os.path.join(ROOT_DIR, "docs", "DB-Schema-Overview.docx")
    doc.save(out_path)
    print(f"\n  Saved: {out_path}")


if __name__ == "__main__":
    main()
