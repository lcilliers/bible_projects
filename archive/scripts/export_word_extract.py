"""
export_word_extract.py
══════════════════════
Generic full-extract for any word in the word registry.

Usage:
    python scripts/export_word_extract.py --word love
    python scripts/export_word_extract.py --word grief
    python scripts/export_word_extract.py          # defaults to --word love

Chain extracted:
    word_registry
      → wa_file_index           (all parts)
      → mti_terms               (+ mti_term_flags, mti_term_cross_refs)
      → wa_term_inventory       (+ related_words, root_family, phase2_flags)
      → wa_cross_registry_links
      → wa_data_quality_flags
      → wa_verse_records        (all verses per term)

Output:
    outputs/docx/{WORD}-full-extract-YYYY-MM-DD.docx
"""

import argparse
import datetime
import os
import sys

ROOT_DIR = os.path.join(os.path.dirname(__file__), "..")
sys.path.insert(0, ROOT_DIR)

from analytics.db_client import get_connection  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt, RGBColor  # noqa: E402


# ── Utility helpers ────────────────────────────────────────────────────────────

def _safe(v) -> str:
    """Coerce any value to a printable string."""
    return "" if v is None else str(v)


def _bool_str(v) -> str:
    if v is None:
        return ""
    return "Yes" if int(v) else "No"


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ── Document helpers ───────────────────────────────────────────────────────────

def _kv_table(doc, data: list[tuple], col_widths=(5, 11)):
    """Two-column key-value table."""
    tbl = doc.add_table(rows=len(data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(data):
        kc = tbl.rows[i].cells[0]
        kc.width = Cm(col_widths[0])
        kr = kc.paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(9)
        _set_cell_bg(kc, "DCE6F1")

        vc = tbl.rows[i].cells[1]
        vc.width = Cm(col_widths[1])
        vr = vc.paragraphs[0].add_run(_safe(v))
        vr.font.size = Pt(9)
    return tbl


def _data_table(doc, headers: list, rows: list,
                header_color="1F497D", col_widths=None):
    """Multi-column data table with styled header."""
    if not rows:
        doc.add_paragraph("(no data)").runs[0].font.size = Pt(9)
        return

    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, header_color)
        if col_widths:
            cell.width = Cm(col_widths[j])

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            r = cell.paragraphs[0].add_run(_safe(val))
            r.font.size = Pt(8)
            if col_widths:
                cell.width = Cm(col_widths[j])
            if i % 2 == 0:
                _set_cell_bg(cell, "EEF3F8")
    return tbl


def _divider(doc):
    doc.add_paragraph("─" * 80).runs[0].font.size = Pt(7)


# ── Data queries ───────────────────────────────────────────────────────────────

def _placeholders(ids: list) -> str:
    return ",".join("?" * len(ids))


def _fetch_all(conn, sql: str, params=()) -> list[dict]:
    return [dict(r) for r in conn.execute(sql, params).fetchall()]


def load_word_data(conn, word: str) -> dict:
    """
    Run all queries for *word* and return a single dict of datasets.

    Query strategy
    ──────────────
    • Every query hits a dedicated index — no full-table scans.
    • wa_verse_records: ORDER BY is intentionally omitted from SQL.
      SQLite cannot use the composite index idx_wavr_file_term_pos when
      file_id is supplied as an IN list (it loops per value and must
      merge streams, forcing a temp B-tree).  Instead we fetch unordered
      and sort in Python — trivially fast for typical result sizes.
    • For single-part words (one file_id) the composite index IS used
      end-to-end (SEARCH ... idx_wavr_file_term_pos, no temp B-tree).
    """
    # 1. word_registry
    reg_rows = _fetch_all(conn,
        "SELECT * FROM word_registry WHERE word = ? COLLATE NOCASE", (word,))
    if not reg_rows:
        return {}
    reg = reg_rows[0]
    registry_id = str(reg["id"])

    # 2. wa_file_index  — indexed on word_registry_fk (idx_wa_fi_wrfk)
    files = _fetch_all(conn,
        "SELECT * FROM wa_file_index "
        "WHERE word_registry_fk = ? ORDER BY part_number",
        (reg["id"],))
    file_ids = [r["id"] for r in files]

    # 3. mti_terms  — indexed on owning_registry (idx_mti_terms_registry)
    mti = _fetch_all(conn,
        "SELECT * FROM mti_terms "
        "WHERE owning_registry = ? ORDER BY language, transliteration",
        (registry_id,))
    mti_ids = [r["id"] for r in mti]

    # 3a. mti_term_flags  — covering index on (mti_term_id, flag_id)
    mti_flags: dict[int, list[str]] = {}
    if mti_ids:
        ph = _placeholders(mti_ids)
        for r in conn.execute(
            f"SELECT f.mti_term_id, p.flag_code, p.description "
            f"FROM mti_term_flags f "
            f"JOIN phase2_flag_types p ON p.id = f.flag_id "
            f"WHERE f.mti_term_id IN ({ph})", mti_ids
        ).fetchall():
            mti_flags.setdefault(r[0], []).append(
                r[1] + (" — " + r[2] if r[2] else ""))

    # 3b. mti_term_cross_refs  — indexed on mti_term_id (idx_cross_refs_term_id)
    mti_xrefs: dict[int, list[dict]] = {}
    if mti_ids:
        ph = _placeholders(mti_ids)
        for r in _fetch_all(conn,
            f"SELECT * FROM mti_term_cross_refs WHERE mti_term_id IN ({ph})", mti_ids):
            mti_xrefs.setdefault(r["mti_term_id"], []).append(r)

    # 4. wa_term_inventory  — indexed on file_id (idx_wa_ti_file)
    inv = [] if not file_ids else _fetch_all(conn,
        f"SELECT * FROM wa_term_inventory "
        f"WHERE file_id IN ({_placeholders(file_ids)}) "
        f"ORDER BY file_id, language, transliteration",
        file_ids)
    inv_ids = [r["id"] for r in inv]

    # 4a. wa_term_related_words  — indexed on term_inv_id (idx_wa_rw)
    related: dict[int, list[dict]] = {}
    if inv_ids:
        ph = _placeholders(inv_ids)
        for r in _fetch_all(conn,
            f"SELECT * FROM wa_term_related_words "
            f"WHERE term_inv_id IN ({ph}) ORDER BY term_inv_id, gloss", inv_ids):
            related.setdefault(r["term_inv_id"], []).append(r)

    # 4b. wa_term_root_family  — indexed on term_inv_id (idx_wa_rf)
    roots: dict[int, list[dict]] = {}
    if inv_ids:
        ph = _placeholders(inv_ids)
        for r in _fetch_all(conn,
            f"SELECT * FROM wa_term_root_family "
            f"WHERE term_inv_id IN ({ph}) ORDER BY term_inv_id, root_code", inv_ids):
            roots.setdefault(r["term_inv_id"], []).append(r)

    # 4c. wa_term_phase2_flags  — covering index (term_inv_id, flag_id)
    p2flags: dict[int, list[str]] = {}
    if inv_ids:
        ph = _placeholders(inv_ids)
        for r in conn.execute(
            f"SELECT f.term_inv_id, t.flag_code, t.description "
            f"FROM wa_term_phase2_flags f "
            f"JOIN phase2_flag_types t ON t.id = f.flag_id "
            f"WHERE f.term_inv_id IN ({ph})", inv_ids
        ).fetchall():
            p2flags.setdefault(r[0], []).append(
                r[1] + (" — " + r[2] if r[2] else ""))

    # 5. wa_cross_registry_links  — indexed on file_id (idx_wa_xrl)
    crl = [] if not file_ids else _fetch_all(conn,
        f"SELECT crl.id, crl.file_id, crl.linked_word, crl.linked_registry_id, "
        f"ct.type_code AS connection_type, crl.connecting_term, crl.note, crl.last_changed "
        f"FROM wa_cross_registry_links crl "
        f"JOIN wa_crosslink_type ct ON ct.id = crl.connection_type_id "
        f"WHERE crl.file_id IN ({_placeholders(file_ids)}) "
        f"ORDER BY crl.file_id, crl.linked_word",
        file_ids)

    # 6. wa_data_quality_flags  — indexed on file_id, joins wa_quality_flag_types
    dqf = [] if not file_ids else _fetch_all(conn,
        f"SELECT dq.id, dq.file_id, dq.term_id, dq.flag_id, "
        f"qt.flag_group, qt.flag_code, dq.description, dq.last_changed "
        f"FROM wa_data_quality_flags dq "
        f"JOIN wa_quality_flag_types qt ON qt.id = dq.flag_id "
        f"WHERE dq.file_id IN ({_placeholders(file_ids)}) "
        f"ORDER BY dq.file_id, dq.term_id",
        file_ids)

    # 7. wa_verse_records  — indexed on file_id (idx_wa_vr_file).
    #    ORDER BY removed from SQL; sorted in Python below.
    #    (See docstring for idx_wavr_file_term_pos note.)
    verses_raw = [] if not file_ids else _fetch_all(conn,
        f"SELECT * FROM wa_verse_records "
        f"WHERE file_id IN ({_placeholders(file_ids)})",
        file_ids)

    # Sort in Python: group by term, then canonical Bible order
    verses_raw.sort(key=lambda r: (
        _safe(r["term_id"]),
        r["book_id"] or 0,
        r["chapter"] or 0,
        r["verse_num"] or 0,
    ))
    verses: dict[str, list[dict]] = {}
    for r in verses_raw:
        verses.setdefault(_safe(r["term_id"]), []).append(r)

    # Reference tables
    flag_types = _fetch_all(conn,
        "SELECT * FROM phase2_flag_types ORDER BY id")
    quality_flag_types = _fetch_all(conn,
        "SELECT * FROM wa_quality_flag_types ORDER BY flag_group, id")

    return dict(
        reg=reg, files=files, file_ids=file_ids,
        mti=mti, mti_flags=mti_flags, mti_xrefs=mti_xrefs,
        inv=inv, related=related, roots=roots, p2flags=p2flags,
        crl=crl, dqf=dqf, verses=verses,
        flag_types=flag_types,
        quality_flag_types=quality_flag_types,
    )


# ── Document builder ───────────────────────────────────────────────────────────

def build_document(data: dict, word: str) -> Document:
    reg       = data["reg"]
    files     = data["files"]
    mti       = data["mti"]
    mti_flags = data["mti_flags"]
    mti_xrefs = data["mti_xrefs"]
    inv       = data["inv"]
    related   = data["related"]
    roots     = data["roots"]
    p2flags   = data["p2flags"]
    crl       = data["crl"]
    dqf       = data["dqf"]
    verses    = data["verses"]
    flag_types= data["flag_types"]
    quality_flag_types = data["quality_flag_types"]

    total_verses = sum(len(v) for v in verses.values())

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2)

    # ── Title ──────────────────────────────────────────────────────────────
    tp = doc.add_heading(f"{word.upper()} — Comprehensive Word Study Extract", 0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Registry #{reg['id']}  |  Generated: {datetime.date.today().isoformat()}"
        f"  |  Terms: {len(inv)}  |  Verses: {total_verses}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S1 — WORD REGISTRY
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading("1. Word Registry  (word_registry)", 1)
    _kv_table(doc, [(k, reg[k]) for k in reg.keys()])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S2 — WA FILE INDEX
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(f"2. WA File Index  (wa_file_index)  — {len(files)} parts", 1)
    fi_cols = [
        "id","filename","registry_id","word_registry_fk","word",
        "part_number","total_parts","is_split","schema_version","phase",
        "produced_date","source_file","translation_used","specification",
        "revision_note","source_list","category","testament_coverage",
        "root_families_in_prior_parts","last_changed",
    ]
    _data_table(doc, fi_cols,
                [[r.get(c) for c in fi_cols] for r in files])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S3 — MTI TERMS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(f"3. MTI Terms  (mti_terms)  — {len(mti)} terms", 1)

    for n, m in enumerate(mti, 1):
        doc.add_heading(
            f"3.{n}  {m['transliteration']}  ({m['strongs_number']})  —  {m['gloss']}",
            2)
        _kv_table(doc, [(k, m[k]) for k in m.keys()])

        flags = mti_flags.get(m["id"], [])
        if flags:
            doc.add_paragraph("Flags: " + "; ".join(flags)
                               ).runs[0].font.size = Pt(8)

        xrefs = mti_xrefs.get(m["id"], [])
        if xrefs:
            p = doc.add_paragraph()
            p.add_run("MTI Cross-References:").bold = True
            p.runs[0].font.size = Pt(9)
            _data_table(doc,
                ["id","term_id","registry","word","part","word_data_reference"],
                [[x[c] for c in ["id","term_id","registry","word","part","word_data_reference"]]
                 for x in xrefs],
                header_color="4472C4")
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S4 — WA TERM INVENTORY  (per-term sub-sections)
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(
        f"4. WA Term Inventory  (wa_term_inventory)  — {len(inv)} terms", 1)
    doc.add_paragraph(
        "Each entry includes related words, root family, phase-2 flags, and all associated verses."
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for n, t in enumerate(inv, 1):
        tid = t.get("term_id") or t.get("strongs_number") or ""
        doc.add_heading(
            f"4.{n}  {t['transliteration']}  ({t['strongs_number'] or t['term_id']})"
            f"  —  {t['step_search_gloss']}  [{t['language']}]",
            2)

        # Core row
        inv_fields = [
            "id","file_id","language","term_id","strongs_number","transliteration",
            "step_search_gloss","word_analysis_gloss",
            "occurrence_count","occurrence_count_qualifier","testament",
            "god_as_subject","somatic_link","causative_form_present",
            "also_spelled","status_note","last_changed",
        ]
        _kv_table(doc, [(f, _bool_str(t[f]) if f in
                         ("god_as_subject","somatic_link","causative_form_present")
                         else t.get(f))
                        for f in inv_fields])

        for long_field, label in (("meaning","meaning"),
                                  ("meaning_numbered","meaning_numbered"),
                                  ("lsj_entry","lsj_entry")):
            if t.get(long_field):
                p = doc.add_paragraph()
                p.add_run(f"{label}:  ").bold = True
                p.runs[0].font.size = Pt(9)
                p.add_run(_safe(t[long_field])).font.size = Pt(8 if long_field == "lsj_entry" else 9)

        doc.add_paragraph()

        # 4.n.a  Related Words
        rw = related.get(t["id"], [])
        doc.add_heading(f"4.{n}.a  Related Words  ({len(rw)} entries)", 3)
        _data_table(doc,
            ["id","term_inv_id","gloss","transliteration"],
            [[r["id"],r["term_inv_id"],r["gloss"],r["transliteration"]] for r in rw],
            header_color="375623",
            col_widths=[2, 2, 7, 5])
        doc.add_paragraph()

        # 4.n.b  Root Family
        rf = roots.get(t["id"], [])
        doc.add_heading(f"4.{n}.b  Root Family  ({len(rf)} entries)", 3)
        _data_table(doc,
            ["id","term_inv_id","root_code"],
            [[r["id"],r["term_inv_id"],r["root_code"]] for r in rf],
            header_color="7030A0",
            col_widths=[2, 2, 12])
        doc.add_paragraph()

        # 4.n.c  Phase-2 Flags
        doc.add_heading(f"4.{n}.c  Phase-2 Flags", 3)
        inline_flags = []
        if t.get("god_as_subject"):
            inline_flags.append("GOD_AS_SUBJECT (inline boolean)")
        if t.get("somatic_link"):
            inline_flags.append("SOMATIC_INNER_LINK (inline boolean)")
        if t.get("causative_form_present"):
            inline_flags.append("CAUSATIVE_OF_INNER_STATE (inline boolean)")
        all_flags = inline_flags + p2flags.get(t["id"], [])
        if all_flags:
            for f in all_flags:
                doc.add_paragraph(f"• {f}", style="List Bullet"
                                  ).runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph("(none)").runs[0].font.size = Pt(9)
        doc.add_paragraph()

        # 4.n.d  Verse Records
        vv = verses.get(tid, [])
        doc.add_heading(f"4.{n}.d  Verse Records  ({len(vv)} verses)", 3)
        _data_table(doc,
            ["id","file_id","term_id","transliteration",
             "testament","reference","chapter","verse_num",
             "translation","verse_text","note","last_changed"],
            [[v["id"],v["file_id"],v["term_id"],v["transliteration"],
              v["testament"],v["reference"],v["chapter"],v["verse_num"],
              v["translation"],v["verse_text"],v["note"],v["last_changed"]]
             for v in vv],
            header_color="C00000",
            col_widths=[1.5,1.5,2.5,2.5,1.5,2,3,1.2,1.2,1.5,8,5,3.5])
        doc.add_paragraph()
        _divider(doc)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S5 — CROSS REGISTRY LINKS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(
        f"5. Cross Registry Links  (wa_cross_registry_links)  — {len(crl)} links", 1)
    _data_table(doc,
        ["id","file_id","linked_word","linked_registry_id",
         "connection_type","connecting_term","note","last_changed"],
        [[r["id"],r["file_id"],r["linked_word"],r["linked_registry_id"],
          r["connection_type"],r["connecting_term"],r["note"],r["last_changed"]]
         for r in crl],
        header_color="833C00",
        col_widths=[1.5,1.5,3,2.5,3.5,3,8,3])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # S6 — DATA QUALITY FLAGS
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(
        f"6. Data Quality Flags  (wa_data_quality_flags)  — {len(dqf)} flags", 1)
    _data_table(doc,
        ["id","file_id","term_id","flag_group","flag_code","description","last_changed"],
        [[r["id"],r["file_id"],r["term_id"],r["flag_group"],r["flag_code"],r["description"],r["last_changed"]]
         for r in dqf],
        header_color="4472C4",
        col_widths=[1.5,1.5,3,3,4,10,3])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # Appendix A — Phase-2 Flag Types reference
    # ═══════════════════════════════════════════════════════════════════════
    doc.add_heading(
        "Appendix A — Phase-2 Flag Types  (phase2_flag_types)", 1)
    _data_table(doc,
        ["id","flag_code","description"],
        [[r["id"],r["flag_code"],r["description"]] for r in flag_types],
        header_color="595959",
        col_widths=[1.5,6,10])
    doc.add_paragraph()

    doc.add_heading(
        "Appendix B — Data Quality Flag Types  (wa_quality_flag_types)", 1)
    _data_table(doc,
        ["id","flag_group","flag_code","description"],
        [[r["id"],r["flag_group"],r["flag_code"],r["description"]] for r in quality_flag_types],
        header_color="4472C4",
        col_widths=[1.5,3.5,5,10])

    return doc


# ── Entry point ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Export a full word-study extract for any registry word.")
    parser.add_argument(
        "--word", default="love",
        help="Word to extract (must exist in word_registry). Default: love")
    args = parser.parse_args()
    word = args.word.lower().strip()

    conn = get_connection()
    print(f"Loading data for '{word}' …")
    data = load_word_data(conn, word)
    conn.close()

    if not data:
        print(f"ERROR: '{word}' not found in word_registry.")
        sys.exit(1)

    reg       = data["reg"]
    files     = data["files"]
    mti       = data["mti"]
    inv       = data["inv"]
    related   = data["related"]
    roots     = data["roots"]
    crl       = data["crl"]
    dqf       = data["dqf"]
    verses    = data["verses"]

    print(f"  Registry #{reg['id']}  |  {len(files)} parts  |  "
          f"{len(mti)} MTI terms  |  {len(inv)} inventory terms  |  "
          f"{sum(len(v) for v in verses.values())} verses")

    print("Building document …")
    doc = build_document(data, word)

    out_dir = os.path.join(ROOT_DIR, "outputs", "docx")
    os.makedirs(out_dir, exist_ok=True)
    date_str = datetime.date.today().isoformat()
    out_path = os.path.join(out_dir, f"{word.upper()}-full-extract-{date_str}.docx")
    doc.save(out_path)

    print(f"\n  Saved: {out_path}")
    print(f"  word_registry:         1 row")
    print(f"  wa_file_index:         {len(files)} rows")
    print(f"  mti_terms:             {len(mti)} rows")
    print(f"  wa_term_inventory:     {len(inv)} rows")
    print(f"  wa_term_related_words: {sum(len(v) for v in related.values())} rows")
    print(f"  wa_term_root_family:   {sum(len(v) for v in roots.values())} rows")
    print(f"  wa_cross_reg_links:    {len(crl)} rows")
    print(f"  wa_data_quality_flags: {len(dqf)} rows")
    print(f"  wa_verse_records:      {sum(len(v) for v in verses.values())} rows")


if __name__ == "__main__":
    main()
