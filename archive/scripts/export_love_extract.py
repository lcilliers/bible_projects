"""
export_love_extract.py
══════════════════════
Full extract of all database tables related to the word "love" (registry #103).

Chain: word_registry → wa_file_index → mti_terms → wa_term_inventory →
       wa_term_related_words, wa_term_root_family, wa_term_phase2_flags →
       wa_cross_registry_links, wa_data_quality_flags → wa_verse_records

Output: outputs/docx/LOVE-full-extract-YYYY-MM-DD.docx
"""

import os
import sys
import datetime

ROOT_DIR = os.path.join(os.path.dirname(__file__), "..")
sys.path.insert(0, ROOT_DIR)

from analytics.db_client import get_connection

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Helpers ──────────────────────────────────────────────────────────────────

def _safe(v):
    """Convert any value to a printable string, handling None and encoding issues."""
    if v is None:
        return ""
    return str(v)


def _bool_str(v):
    if v is None:
        return ""
    return "Yes" if int(v) else "No"


def set_cell_bg(cell, hex_color: str):
    """Set background fill colour on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_kv_table(doc, data: list[tuple], col_widths=(5, 11)):
    """Add a 2-column key-value table."""
    tbl = doc.add_table(rows=len(data), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(data):
        # Key cell
        kc = tbl.rows[i].cells[0]
        kc.width = Cm(col_widths[0])
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(9)
        set_cell_bg(kc, "DCE6F1")

        # Value cell
        vc = tbl.rows[i].cells[1]
        vc.width = Cm(col_widths[1])
        vp = vc.paragraphs[0]
        vr = vp.add_run(_safe(v))
        vr.font.size = Pt(9)
    return tbl


def add_data_table(doc, headers: list, rows: list, header_color="1F497D", col_widths=None):
    """Add a multi-column data table with a styled header row."""
    if not rows:
        doc.add_paragraph("(no data)")
        return
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_color)
        if col_widths:
            cell.width = Cm(col_widths[j])

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.rows[i + 1].cells[j]
            p = cell.paragraphs[0]
            r = p.add_run(_safe(val))
            r.font.size = Pt(8)
            if col_widths:
                cell.width = Cm(col_widths[j])
            if i % 2 == 0:
                set_cell_bg(cell, "EEF3F8")
    return tbl


def add_section_divider(doc):
    doc.add_paragraph("─" * 80).runs[0].font.size = Pt(7)


# ── Main extraction ───────────────────────────────────────────────────────────

def run():
    conn = get_connection()

    # ── Resolve love registry ──────────────────────────────────────────────
    reg = conn.execute(
        "SELECT * FROM word_registry WHERE word = 'love'"
    ).fetchone()
    if reg is None:
        print("ERROR: 'love' not found in word_registry.")
        conn.close()
        return

    reg = dict(reg)
    registry_id = str(reg["id"])  # '103'

    # ── wa_file_index ──────────────────────────────────────────────────────
    file_rows = conn.execute(
        "SELECT * FROM wa_file_index WHERE word_registry_fk = ? ORDER BY part_number",
        (reg["id"],),
    ).fetchall()
    file_rows = [dict(r) for r in file_rows]
    file_ids = [r["id"] for r in file_rows]

    # ── mti_terms ──────────────────────────────────────────────────────────
    if file_ids:
        mti_rows = conn.execute(
            f"SELECT * FROM mti_terms WHERE owning_registry = ? ORDER BY language, transliteration",
            (registry_id,),
        ).fetchall()
        mti_rows = [dict(r) for r in mti_rows]
    else:
        mti_rows = []

    mti_ids = [r["id"] for r in mti_rows]

    # ── mti_term_flags ──────────────────────────────────────────────────────
    mti_flags_map = {}  # term_id → [flag descriptions]
    if mti_ids:
        placeholders = ",".join("?" * len(mti_ids))
        mti_flag_rows = conn.execute(
            f"SELECT mtf.mti_term_id, mpf.flag_code, mpf.description "
            f"FROM mti_term_flags mtf "
            f"JOIN phase2_flag_types mpf ON mpf.id = mtf.flag_id "
            f"WHERE mtf.mti_term_id IN ({placeholders})",
            mti_ids,
        ).fetchall()
        for r in mti_flag_rows:
            mti_flags_map.setdefault(r[0], []).append(
                f"{r[1]}" + (f" — {r[2]}" if r[2] else "")
            )

    # ── mti_term_cross_refs ─────────────────────────────────────────────────
    mti_xref_map = {}  # mti_term_id → [dict]
    if mti_ids:
        placeholders = ",".join("?" * len(mti_ids))
        mti_xref_rows = conn.execute(
            f"SELECT * FROM mti_term_cross_refs WHERE mti_term_id IN ({placeholders})",
            mti_ids,
        ).fetchall()
        for r in mti_xref_rows:
            mti_xref_map.setdefault(r["mti_term_id"], []).append(dict(r))

    # ── wa_term_inventory ──────────────────────────────────────────────────
    if file_ids:
        placeholders = ",".join("?" * len(file_ids))
        inv_rows = conn.execute(
            f"SELECT * FROM wa_term_inventory WHERE file_id IN ({placeholders}) "
            f"ORDER BY file_id, language, transliteration",
            file_ids,
        ).fetchall()
        inv_rows = [dict(r) for r in inv_rows]
    else:
        inv_rows = []

    inv_ids = [r["id"] for r in inv_rows]

    # ── wa_term_related_words ───────────────────────────────────────────────
    related_map = {}  # term_id → [dict]
    if inv_ids:
        placeholders = ",".join("?" * len(inv_ids))
        rw_rows = conn.execute(
            f"SELECT * FROM wa_term_related_words WHERE term_inv_id IN ({placeholders}) "
            f"ORDER BY term_inv_id, gloss",
            inv_ids,
        ).fetchall()
        for r in rw_rows:
            related_map.setdefault(r["term_inv_id"], []).append(dict(r))

    # ── wa_term_root_family ─────────────────────────────────────────────────
    root_map = {}  # term_id → [dict]
    if inv_ids:
        placeholders = ",".join("?" * len(inv_ids))
        rf_rows = conn.execute(
            f"SELECT * FROM wa_term_root_family WHERE term_inv_id IN ({placeholders}) "
            f"ORDER BY term_inv_id, root_code",
            inv_ids,
        ).fetchall()
        for r in rf_rows:
            root_map.setdefault(r["term_inv_id"], []).append(dict(r))

    # ── wa_term_phase2_flags ────────────────────────────────────────────────
    p2flag_map = {}  # term_id → [flag descriptions]
    if inv_ids:
        placeholders = ",".join("?" * len(inv_ids))
        p2f_rows = conn.execute(
            f"SELECT wtpf.term_inv_id, wft.flag_code, wft.description "
            f"FROM wa_term_phase2_flags wtpf "
            f"JOIN phase2_flag_types wft ON wft.id = wtpf.flag_id "
            f"WHERE wtpf.term_inv_id IN ({placeholders})",
            inv_ids,
        ).fetchall()
        for r in p2f_rows:
            p2flag_map.setdefault(r[0], []).append(
                f"{r[1]}" + (f" — {r[2]}" if r[2] else "")
            )

    # ── phase2_flag_types reference ──────────────────────────────────────────
    flag_types = conn.execute("SELECT * FROM phase2_flag_types ORDER BY id").fetchall()
    flag_types = [dict(r) for r in flag_types]
    quality_flag_types = conn.execute(
        "SELECT * FROM wa_quality_flag_types ORDER BY flag_group, id").fetchall()
    quality_flag_types = [dict(r) for r in quality_flag_types]

    # ── wa_cross_registry_links ─────────────────────────────────────────────
    if file_ids:
        placeholders = ",".join("?" * len(file_ids))
        crl_rows = conn.execute(
            f"SELECT crl.id, crl.file_id, crl.linked_word, crl.linked_registry_id, "
            f"ct.type_code AS connection_type, crl.connecting_term, crl.note, crl.last_changed "
            f"FROM wa_cross_registry_links crl "
            f"JOIN wa_crosslink_type ct ON ct.id = crl.connection_type_id "
            f"WHERE crl.file_id IN ({placeholders}) "
            f"ORDER BY crl.file_id, crl.linked_word",
            file_ids,
        ).fetchall()
        crl_rows = [dict(r) for r in crl_rows]
    else:
        crl_rows = []

    # ── wa_data_quality_flags ───────────────────────────────────────────────
    if file_ids:
        placeholders = ",".join("?" * len(file_ids))
        dqf_rows = conn.execute(
            f"SELECT dq.id, dq.file_id, dq.term_id, dq.flag_id, "
            f"qt.flag_group, qt.flag_code, dq.description, dq.last_changed "
            f"FROM wa_data_quality_flags dq "
            f"JOIN wa_quality_flag_types qt ON qt.id = dq.flag_id "
            f"WHERE dq.file_id IN ({placeholders}) "
            f"ORDER BY dq.file_id, dq.term_id",
            file_ids,
        ).fetchall()
        dqf_rows = [dict(r) for r in dqf_rows]
    else:
        dqf_rows = []

    # ── wa_verse_records ────────────────────────────────────────────────────
    # Map term_id → [verse records]
    verse_map = {}  # term_id → [dict]
    if file_ids:
        placeholders = ",".join("?" * len(file_ids))
        vr_rows = conn.execute(
            f"SELECT * FROM wa_verse_records "
            f"WHERE file_id IN ({placeholders}) "
            f"ORDER BY term_id, book_id, chapter, verse_num",
            file_ids,
        ).fetchall()
        for r in vr_rows:
            rd = dict(r)
            verse_map.setdefault(rd["term_id"], []).append(rd)

    conn.close()

    # ═════════════════════════════════════════════════════════════════════════
    # Build DOCX
    # ═════════════════════════════════════════════════════════════════════════
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ── Title ──────────────────────────────────────────────────────────────
    title_p = doc.add_heading("LOVE — Comprehensive Word Study Extract", 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph(
        f"Registry #103  |  Generated: {datetime.date.today().isoformat()}  |  "
        f"Terms: {len(inv_rows)}  |  Verses: {sum(len(v) for v in verse_map.values())}"
    )
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.runs[0].font.size = Pt(9)
    meta_p.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1 — WORD REGISTRY
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. Word Registry", 1)
    doc.add_paragraph(
        "The authoritative entry for 'love' in the word registry. "
        "All subsequent data traces back to this record."
    ).runs[0].font.size = Pt(9)

    add_kv_table(doc, [
        ("id",                   reg["id"]),
        ("no",                   reg["no"]),
        ("word",                 reg["word"]),
        ("source_list",          reg["source_list"]),
        ("category_hint",        reg["category_hint"]),
        ("phase1_input_file",    reg["phase1_input_file"]),
        ("phase1_status",        reg["phase1_status"]),
        ("phase1_output_file",   reg["phase1_output_file"]),
        ("phase2_datasets",      reg["phase2_datasets"]),
        ("notes",                reg["notes"]),
    ])
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2 — WA FILE INDEX
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. WA File Index  (wa_file_index)", 1)
    doc.add_paragraph(
        f"Three data-collection parts for registry #103. "
        f"Parts 1–2 cover OT terms; Part 3 covers NT terms."
    ).runs[0].font.size = Pt(9)

    fi_headers = [
        "id", "filename", "registry_id", "word_registry_fk", "word",
        "part_number", "total_parts", "is_split", "schema_version", "phase",
        "produced_date", "source_file", "translation_used", "specification",
        "revision_note", "source_list", "category", "testament_coverage",
        "root_families_in_prior_parts", "last_changed",
    ]
    add_data_table(
        doc,
        fi_headers,
        [[r.get(h) for h in fi_headers] for r in file_rows],
        col_widths=None,
    )
    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 3 — MTI TERMS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. MTI Terms  (mti_terms)", 1)
    doc.add_paragraph(
        f"{len(mti_rows)} terms extracted for registry #103, spanning Hebrew and Greek."
    ).runs[0].font.size = Pt(9)

    for mti in mti_rows:
        add_heading(doc, f"3.{mti_rows.index(mti)+1}  {mti['transliteration']}  ({mti['strongs_number']})  —  {mti['gloss']}", 2)

        add_kv_table(doc, [
            ("id",                   mti["id"]),
            ("strongs_number",       mti["strongs_number"]),
            ("transliteration",      mti["transliteration"]),
            ("gloss",                mti["gloss"]),
            ("language",             mti["language"]),
            ("owning_registry",      mti["owning_registry"]),
            ("owning_word",          mti["owning_word"]),
            ("owning_part",          mti["owning_part"]),
            ("word_data_reference",  mti["word_data_reference"]),
            ("status",               mti["status"]),
            ("status_note",          mti["status_note"]),
            ("exclusion_reason",     mti["exclusion_reason"]),
            ("extraction_date",      mti["extraction_date"]),
            ("strongs_reconciled",   _bool_str(mti["strongs_reconciled"])),
            ("anchor_note",          mti["anchor_note"]),
            ("last_changed",         mti["last_changed"]),
        ])

        # MTI term flags
        flags = mti_flags_map.get(mti["id"], [])
        if flags:
            doc.add_paragraph("MTI Flags: " + "; ".join(flags)).runs[0].font.size = Pt(8)

        # MTI cross-refs
        xrefs = mti_xref_map.get(mti["id"], [])
        if xrefs:
            p = doc.add_paragraph()
            p.add_run("MTI Cross-References:").bold = True
            p.runs[0].font.size = Pt(9)
            add_data_table(
                doc,
                ["id", "term_id", "registry", "word", "part", "word_data_reference"],
                [[x["id"], x["term_id"], x["registry"], x["word"], x["part"], x["word_data_reference"]] for x in xrefs],
                header_color="4472C4",
            )

        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 4 — WA TERM INVENTORY (with sub-sections)
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. WA Term Inventory  (wa_term_inventory)", 1)
    doc.add_paragraph(
        f"{len(inv_rows)} terms in the inventory for registry #103. "
        "Each entry includes related words, root family, phase-2 flags, and all verse records."
    ).runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for inv in inv_rows:
        term_idx = inv_rows.index(inv) + 1
        add_heading(
            doc,
            f"4.{term_idx}  {inv['transliteration']}  ({inv['strongs_number'] or inv['term_id']})  "
            f"—  {inv['step_search_gloss']}  [{inv['language']}]",
            2,
        )

        # Core inventory row
        add_kv_table(doc, [
            ("id",                        inv["id"]),
            ("file_id",                   inv["file_id"]),
            ("language",                  inv["language"]),
            ("term_id",                   inv["term_id"]),
            ("strongs_number",            inv["strongs_number"]),
            ("transliteration",           inv["transliteration"]),
            ("step_search_gloss",         inv["step_search_gloss"]),
            ("word_analysis_gloss",       inv["word_analysis_gloss"]),
            ("occurrence_count",          inv["occurrence_count"]),
            ("occurrence_count_qualifier",inv["occurrence_count_qualifier"]),
            ("testament",                 inv["testament"]),
            ("god_as_subject",            _bool_str(inv["god_as_subject"])),
            ("somatic_link",              _bool_str(inv["somatic_link"])),
            ("causative_form_present",    _bool_str(inv["causative_form_present"])),
            ("also_spelled",              inv["also_spelled"]),
            ("status_note",               inv["status_note"]),
            ("last_changed",              inv["last_changed"]),
        ])

        # Meaning (possibly long)
        if inv.get("meaning"):
            p = doc.add_paragraph()
            p.add_run("meaning:  ").bold = True
            p.runs[0].font.size = Pt(9)
            p.add_run(_safe(inv["meaning"])).font.size = Pt(9)

        if inv.get("meaning_numbered"):
            doc.add_paragraph(
                f"meaning_numbered: {inv['meaning_numbered']}"
            ).runs[0].font.size = Pt(8)

        if inv.get("lsj_entry"):
            p2 = doc.add_paragraph()
            p2.add_run("lsj_entry:  ").bold = True
            p2.runs[0].font.size = Pt(9)
            p2.add_run(_safe(inv["lsj_entry"])).font.size = Pt(8)

        doc.add_paragraph()

        # ── 4.x.a  Related Words ──────────────────────────────────────────
        rw = related_map.get(inv["id"], [])
        add_heading(doc, f"4.{term_idx}.a  Related Words  ({len(rw)} entries)", 3)
        if rw:
            add_data_table(
                doc,
                ["id", "term_inv_id", "gloss", "transliteration"],
                [[r["id"], r["term_inv_id"], r["gloss"], r["transliteration"]] for r in rw],
                header_color="375623",
                col_widths=[2, 2, 7, 5],
            )
        else:
            doc.add_paragraph("(none)").runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── 4.x.b  Root Family ────────────────────────────────────────────
        rf = root_map.get(inv["id"], [])
        add_heading(doc, f"4.{term_idx}.b  Root Family  ({len(rf)} entries)", 3)
        if rf:
            add_data_table(
                doc,
                ["id", "term_inv_id", "root_code"],
                [[r["id"], r["term_inv_id"], r["root_code"]] for r in rf],
                header_color="7030A0",
                col_widths=[2, 2, 12],
            )
        else:
            doc.add_paragraph("(none)").runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── 4.x.c  Phase-2 Flags (wa_term_phase2_flags) ──────────────────
        p2flags = p2flag_map.get(inv["id"], [])
        add_heading(doc, f"4.{term_idx}.c  Phase-2 Flags", 3)
        # Also show inline boolean flags from inventory
        inline_flags = []
        if inv.get("god_as_subject"):
            inline_flags.append("GOD_AS_SUBJECT (inline)")
        if inv.get("somatic_link"):
            inline_flags.append("SOMATIC_INNER_LINK (inline)")
        if inv.get("causative_form_present"):
            inline_flags.append("CAUSATIVE_OF_INNER_STATE (inline)")
        all_flags = inline_flags + p2flags
        if all_flags:
            for f in all_flags:
                doc.add_paragraph(f"• {f}", style="List Bullet").runs[0].font.size = Pt(9)
        else:
            doc.add_paragraph("(none)").runs[0].font.size = Pt(9)

        doc.add_paragraph()

        # ── 4.x.d  Verse Records ──────────────────────────────────────────
        tid = inv.get("term_id") or inv.get("strongs_number")
        verses = verse_map.get(tid, [])
        add_heading(doc, f"4.{term_idx}.d  Verse Records  ({len(verses)} verses)", 3)
        if verses:
            add_data_table(
                doc,
                [
                    "id", "file_id", "term_id", "transliteration",
                    "testament", "reference", "chapter", "verse_num",
                    "translation", "verse_text", "note", "last_changed",
                ],
                [
                    [
                        v["id"], v["file_id"], v["term_id"], v["transliteration"],
                        v["testament"], v["reference"], v["chapter"], v["verse_num"],
                        v["translation"], v["verse_text"], v["note"], v["last_changed"],
                    ]
                    for v in verses
                ],
                header_color="C00000",
                col_widths=[1.5, 1.5, 2.5, 2.5, 1.5, 2, 3, 1.2, 1.2, 1.5, 8, 5, 3.5],
            )
        else:
            doc.add_paragraph("(no verse records)").runs[0].font.size = Pt(9)

        doc.add_paragraph()
        add_section_divider(doc)
        doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 5 — CROSS REGISTRY LINKS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. Cross Registry Links  (wa_cross_registry_links)", 1)
    doc.add_paragraph(
        f"{len(crl_rows)} cross-registry connections found for registry #103."
    ).runs[0].font.size = Pt(9)

    if crl_rows:
        add_data_table(
            doc,
            ["id", "file_id", "linked_word", "linked_registry_id",
             "connection_type", "connecting_term", "note", "last_changed"],
            [
                [r["id"], r["file_id"], r["linked_word"], r["linked_registry_id"],
                 r["connection_type"], r["connecting_term"], r["note"], r["last_changed"]]
                for r in crl_rows
            ],
            header_color="833C00",
            col_widths=[1.5, 1.5, 3, 2.5, 3.5, 3, 8, 3],
        )
    else:
        doc.add_paragraph("(no cross-registry links)").runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 6 — DATA QUALITY FLAGS
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. Data Quality Flags  (wa_data_quality_flags)", 1)
    doc.add_paragraph(
        f"{len(dqf_rows)} data quality flags for registry #103."
    ).runs[0].font.size = Pt(9)

    if dqf_rows:
        add_data_table(
            doc,
            ["id", "file_id", "term_id", "flag_group", "flag", "description", "last_changed"],
            [
                [r["id"], r["file_id"], r["term_id"], r["flag_group"], r["flag_code"], r["description"], r["last_changed"]]
                for r in dqf_rows
            ],
            header_color="4472C4",
            col_widths=[1.5, 1.5, 3, 3, 4, 10, 3],
        )
    else:
        doc.add_paragraph("(no data quality flags)").runs[0].font.size = Pt(9)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # APPENDIX — WA Phase-2 Flag Types Reference
    # ═══════════════════════════════════════════════════════════════════════
    add_heading(doc, "Appendix A — Phase-2 Flag Types  (phase2_flag_types)", 1)
    add_data_table(
        doc,
        ["id", "flag_code", "description"],
        [[r["id"], r["flag_code"], r["description"]] for r in flag_types],
        header_color="595959",
        col_widths=[1.5, 6, 10],
    )
    doc.add_paragraph()
    add_heading(doc, "Appendix B — Data Quality Flag Types  (wa_quality_flag_types)", 1)
    add_data_table(
        doc,
        ["id", "flag_group", "flag_code", "description"],
        [[r["id"], r["flag_group"], r["flag_code"], r["description"]] for r in quality_flag_types],
        header_color="4472C4",
        col_widths=[1.5, 3.5, 5, 10],
    )

    # ── Save ───────────────────────────────────────────────────────────────
    out_dir = os.path.join(ROOT_DIR, "outputs", "docx")
    os.makedirs(out_dir, exist_ok=True)
    date_str = datetime.date.today().isoformat()
    out_path = os.path.join(out_dir, f"LOVE-full-extract-{date_str}.docx")
    doc.save(out_path)
    print(f"\n✓ Saved: {out_path}")
    print(f"  Registry:          1 row")
    print(f"  File Index:        {len(file_rows)} rows")
    print(f"  MTI Terms:         {len(mti_rows)} rows")
    print(f"  Term Inventory:    {len(inv_rows)} rows")
    print(f"  Related Words:     {sum(len(v) for v in related_map.values())} rows")
    print(f"  Root Family:       {sum(len(v) for v in root_map.values())} rows")
    print(f"  Phase-2 Flags:     {sum(len(v) for v in p2flag_map.values())} rows")
    print(f"  Cross-Reg Links:   {len(crl_rows)} rows")
    print(f"  Data Quality Flags:{len(dqf_rows)} rows")
    print(f"  Verse Records:     {sum(len(v) for v in verse_map.values())} rows")


if __name__ == "__main__":
    run()
